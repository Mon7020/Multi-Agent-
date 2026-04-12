"""
文档加载器模块
负责加载和分割各种格式的文档（PDF/Word/TXT）
"""

import os
from typing import List, Dict, Any, Optional
import pdfplumber
import docx
from loguru import logger
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from config.settings import settings


class DocumentLoader:
    """文档加载器，支持PDF、Word、TXT格式"""

    def __init__(self):
        """初始化文档加载器"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.rag.chunk_size,
            chunk_overlap=settings.rag.chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", "，", "、"]
        )
        self.logger = logger.bind(name="document_loader")
        self.logger.info("文档加载器初始化完成")

    def load_document(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """加载本地文档（PDF/Word/TXT）

        Args:
            file_path: 文档路径
            metadata: 可选元数据

        Returns:
            分割后的文档列表

        Raises:
            ValueError: 文件不存在或不支持的文件类型
        """
        if not os.path.exists(file_path):
            raise ValueError(f"文档不存在: {file_path}")

        metadata = metadata or {}
        metadata["file_path"] = file_path
        metadata["update_time"] = datetime.now().isoformat()

        # 根据文件类型读取内容
        if file_path.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise ValueError(f"不支持的文件类型: {file_path}")

        # 分割文本
        chunks = self.text_splitter.split_text(text)
        documents = [Document(page_content=chunk, metadata=metadata) for chunk in chunks]

        self.logger.info(f"文档加载完成: {file_path}，切分为{len(documents)}块")
        return documents

    def load_documents_from_directory(self, directory_path: str) -> List[Document]:
        """从目录加载所有支持的文档

        Args:
            directory_path: 目录路径

        Returns:
            所有文档的合并列表
        """
        if not os.path.isdir(directory_path):
            raise ValueError(f"目录不存在: {directory_path}")

        all_documents = []
        supported_extensions = {".pdf", ".docx", ".txt"}

        for filename in os.listdir(directory_path):
            if any(filename.lower().endswith(ext) for ext in supported_extensions):
                file_path = os.path.join(directory_path, filename)
                try:
                    documents = self.load_document(file_path)
                    all_documents.extend(documents)
                    self.logger.info(f"成功加载文件: {filename} ({len(documents)}块)")
                except Exception as e:
                    self.logger.error(f"加载文件 {filename} 失败: {str(e)}")

        self.logger.info(f"目录 {directory_path} 加载完成，共 {len(all_documents)} 个文档块")
        return all_documents

    def validate_file(self, file_path: str) -> bool:
        """验证文件是否支持

        Args:
            file_path: 文件路径

        Returns:
            是否支持该文件类型
        """
        if not os.path.exists(file_path):
            return False

        supported_extensions = {".pdf", ".docx", ".txt"}
        return any(file_path.lower().endswith(ext) for ext in supported_extensions)


# 导入datetime
from datetime import datetime