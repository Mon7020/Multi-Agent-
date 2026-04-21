"""
CRAG/Self-RAG工具模块（优化版）
使用 ChromaDB 作为向量数据库，轻量级且易于部署
支持 Redis 缓存和 Rerank 重排序
三层架构：Query Understanding -> Retrieval -> Generation Context
"""

import os
import sys
import time
import json
import math
import posixpath
import re
import hashlib
import asyncio
import threading
import zipfile
from typing import List, Dict, Any, Optional, Union, Tuple
from datetime import datetime, timedelta, date
from dataclasses import dataclass, field
from concurrent.futures import ThreadPoolExecutor
from html.parser import HTMLParser
from xml.etree import ElementTree as ET

import pdfplumber
import docx
from pydantic.v1 import BaseModel, Field
from sentence_transformers import SentenceTransformer, CrossEncoder

import chromadb
from chromadb.config import Settings as ChromaSettings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.embeddings import Embeddings
from langchain_core.documents import Document

sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from config.settings import settings
from core.logger import LoggerManager

# 三层架构
from tools.rag.query_understanding import QueryUnderstandingLayer
from tools.rag.generation_context import GenerationContextLayer
from tools.rag.bm25_ranker import ProfessionalBM25
from tools.rag.cache_policy import build_retrieval_cache_key, normalize_retrieval_policy
from tools.rag.chroma_backend import ChromaVectorStoreBackend
from tools.rag.redis_cache_manager import get_cache_manager, RetrievalCache
from tools.rag.vector_store_backend import (
    VectorSearchRequest,
    build_vector_metadata_filter,
    metadata_matches_filter,
)

app_logger = LoggerManager.get_logger("rag_tool")

_XLSX_MAIN_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_XLSX_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_XLSX_DOC_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


class _StructuredHTMLBlockParser(HTMLParser):
    NOISE_TAGS = {"head", "nav", "script", "style", "footer", "noscript"}
    BLOCK_TAGS = {"p", "li"}

    def __init__(self):
        super().__init__()
        self.blocks: List[Dict[str, Any]] = []
        self.heading_stack: List[str] = []
        self._ignored_depth = 0
        self._capture_mode: Optional[str] = None
        self._capture_level: Optional[int] = None
        self._capture_parts: List[str] = []

    @staticmethod
    def _normalize_text(parts: List[str]) -> str:
        text = "".join(parts)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def _flush_heading(self):
        if self._capture_mode != "heading" or self._capture_level is None:
            return
        heading_text = self._normalize_text(self._capture_parts)
        if heading_text:
            level_index = self._capture_level - 1
            self.heading_stack = self.heading_stack[:level_index]
            self.heading_stack.append(heading_text)
        self._capture_mode = None
        self._capture_level = None
        self._capture_parts = []

    def _flush_block(self):
        if self._capture_mode != "block":
            return
        block_text = self._normalize_text(self._capture_parts)
        if block_text:
            self.blocks.append(
                {
                    "heading_path": list(self.heading_stack),
                    "text": block_text,
                }
            )
        self._capture_mode = None
        self._capture_level = None
        self._capture_parts = []

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in self.NOISE_TAGS:
            self._ignored_depth += 1
            return
        if self._ignored_depth:
            return
        if re.fullmatch(r"h[1-6]", tag):
            self._flush_block()
            self._flush_heading()
            self._capture_mode = "heading"
            self._capture_level = int(tag[1])
            self._capture_parts = []
            return
        if tag in self.BLOCK_TAGS:
            self._flush_block()
            self._capture_mode = "block"
            self._capture_level = None
            self._capture_parts = []
            return
        if tag == "br" and self._capture_mode == "block":
            self._capture_parts.append("\n")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self.NOISE_TAGS:
            if self._ignored_depth:
                self._ignored_depth -= 1
            return
        if self._ignored_depth:
            return
        if re.fullmatch(r"h[1-6]", tag):
            self._flush_heading()
            return
        if tag in self.BLOCK_TAGS:
            self._flush_block()

    def handle_data(self, data):
        if self._ignored_depth or self._capture_mode is None:
            return
        self._capture_parts.append(data)

    def close(self):
        super().close()
        self._flush_heading()
        self._flush_block()


@dataclass
class RAGMetrics:
    """RAG系统指标收集"""
    total_requests: int = 0
    successful_requests: int = 0
    failed_requests: int = 0
    total_retrieval_time: float = 0.0
    self_rag_decisions: Dict[str, int] = field(default_factory=lambda: {
        "no_retrieval": 0,
        "retrieval": 0,
        "failed": 0
    })
    documents_processed: int = 0
    cache_hits: int = 0
    cache_misses: int = 0
    avg_retrieval_latency: float = 0.0

    def record_request(self, success: bool, retrieval_time: float):
        self.total_requests += 1
        if success:
            self.successful_requests += 1
        else:
            self.failed_requests += 1
        self.total_retrieval_time += retrieval_time
        if self.total_requests > 0:
            self.avg_retrieval_latency = self.total_retrieval_time / self.total_requests

    def record_self_rag_decision(self, decision_type: str):
        if decision_type in self.self_rag_decisions:
            self.self_rag_decisions[decision_type] += 1

    def record_cache_hit(self):
        self.cache_hits += 1

    def record_cache_miss(self):
        self.cache_misses += 1

    def record_documents_processed(self, count: int):
        self.documents_processed += count

    def get_success_rate(self) -> float:
        if self.total_requests == 0:
            return 0.0
        return self.successful_requests / self.total_requests

    def get_cache_hit_rate(self) -> float:
        total = self.cache_hits + self.cache_misses
        if total == 0:
            return 0.0
        return self.cache_hits / total

    def get_summary(self) -> Dict[str, Any]:
        return {
            "total_requests": self.total_requests,
            "successful_requests": self.successful_requests,
            "failed_requests": self.failed_requests,
            "success_rate": f"{self.get_success_rate():.2%}",
            "avg_retrieval_latency": f"{self.avg_retrieval_latency:.3f}s",
            "self_rag_decisions": self.self_rag_decisions,
            "documents_processed": self.documents_processed,
            "cache_hit_rate": f"{self.get_cache_hit_rate():.2%}",
            "cache_hits": self.cache_hits,
            "cache_misses": self.cache_misses
        }

    def reset(self):
        self.total_requests = 0
        self.successful_requests = 0
        self.failed_requests = 0
        self.total_retrieval_time = 0.0
        self.self_rag_decisions = {"no_retrieval": 0, "retrieval": 0, "failed": 0}
        self.documents_processed = 0
        self.cache_hits = 0
        self.cache_misses = 0
        self.avg_retrieval_latency = 0.0


class LocalCache:
    """简单的内存缓存，支持TTL和LRU淘汰（线程安全版本）"""

    def __init__(self, max_size: int = 1000, default_ttl: int = 3600):
        self._lock = threading.Lock()
        self.cache: Dict[str, Dict[str, Any]] = {}
        self.max_size = max_size
        self.default_ttl = default_ttl
        self.access_order: List[str] = []

    def _generate_key(
        self,
        query: str,
        top_k: int,
        enable_self_rag: bool,
        retrieval_policy: Optional[Dict[str, Any]] = None,
    ) -> str:
        return build_retrieval_cache_key(
            query,
            top_k,
            enable_self_rag,
            retrieval_policy=retrieval_policy,
        )

    def get(
        self,
        query: str,
        top_k: int,
        enable_self_rag: bool,
        retrieval_policy: Optional[Dict[str, Any]] = None,
    ) -> Optional[Dict[str, Any]]:
        key = self._generate_key(query, top_k, enable_self_rag, retrieval_policy)

        with self._lock:
            if key not in self.cache:
                return None

            entry = self.cache[key]
            if datetime.now() > entry['expires_at']:
                self._remove_key_unlocked(key)
                return None

            self._update_access_order_unlocked(key)
            return entry['result']

    def set(
        self,
        query: str,
        top_k: int,
        enable_self_rag: bool,
        result: Dict[str, Any],
        ttl: Optional[int] = None,
        retrieval_policy: Optional[Dict[str, Any]] = None,
    ):
        with self._lock:
            if len(self.cache) >= self.max_size:
                self._evict_lru_unlocked()

            key = self._generate_key(query, top_k, enable_self_rag, retrieval_policy)
            ttl = ttl or self.default_ttl

            self.cache[key] = {
                'result': result,
                'expires_at': datetime.now() + timedelta(seconds=ttl),
                'created_at': datetime.now()
            }
            self._update_access_order_unlocked(key)

    def _update_access_order_unlocked(self, key: str):
        """内部方法，调用时必须持有锁"""
        if key in self.access_order:
            self.access_order.remove(key)
        self.access_order.append(key)

    def _remove_key_unlocked(self, key: str):
        """内部方法，调用时必须持有锁"""
        if key in self.cache:
            del self.cache[key]
        if key in self.access_order:
            self.access_order.remove(key)

    def _evict_lru_unlocked(self):
        """内部方法，调用时必须持有锁"""
        if self.access_order:
            oldest_key = self.access_order[0]
            self._remove_key_unlocked(oldest_key)
            app_logger.debug(f"缓存淘汰: {oldest_key}")

    def clear(self):
        with self._lock:
            self.cache.clear()
            self.access_order.clear()

    def get_stats(self) -> Dict[str, Any]:
        with self._lock:
            return {
                "size": len(self.cache),
                "max_size": self.max_size,
                "default_ttl": self.default_ttl
            }


class RedisCache:
    """Redis缓存实现，支持分布式部署 - 使用JSON序列化替代pickle"""

    def __init__(self, host: str = "localhost", port: int = 6379, db: int = 0,
                 password: str = None, key_prefix: str = "rag:",
                 default_ttl: int = 3600):
        self.key_prefix = key_prefix
        self.default_ttl = default_ttl
        self._redis = None
        self._available = False

        try:
            import redis
            self._redis = redis.Redis(
                host=host,
                port=port,
                db=db,
                password=password,
                decode_responses=False,  # 返回字节以便JSON解析
                socket_connect_timeout=5,
                socket_timeout=5
            )
            # 测试连接
            self._redis.ping()
            self._available = True
            app_logger.info(f"Redis缓存连接成功: {host}:{port}")
        except Exception as e:
            app_logger.warning(f"Redis连接失败，回退到内存缓存: {e}")
            self._redis = None
            self._available = False

    @staticmethod
    def _json_safe_serialize(obj: Any) -> str:
        """将对象转换为JSON安全格式"""
        if isinstance(obj, dict):
            return {k: RedisCache._json_safe_serialize(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [RedisCache._json_safe_serialize(item) for item in obj]
        elif isinstance(obj, (str, int, float, bool, type(None))):
            return obj
        elif isinstance(obj, (datetime, date)):
            return obj.isoformat()
        elif hasattr(obj, '__dict__'):
            # 处理带属性的对象（如Document）
            return str(obj)
        else:
            return str(obj)

    def _generate_key(
        self,
        query: str,
        top_k: int,
        enable_self_rag: bool,
        retrieval_policy: Optional[Dict[str, Any]] = None,
    ) -> str:
        key_hash = build_retrieval_cache_key(
            query,
            top_k,
            enable_self_rag,
            retrieval_policy=retrieval_policy,
        )
        return f"{self.key_prefix}{key_hash}"

    def get(
        self,
        query: str,
        top_k: int,
        enable_self_rag: bool,
        retrieval_policy: Optional[Dict[str, Any]] = None,
    ) -> Optional[Dict[str, Any]]:
        if not self._available or not self._redis:
            return None

        try:
            key = self._generate_key(query, top_k, enable_self_rag, retrieval_policy)
            data = self._redis.get(key)
            if data:
                # 使用JSON反序列化替代pickle
                result = json.loads(data.decode('utf-8'))
                app_logger.debug(f"Redis缓存命中: {key[:20]}...")
                return result
            return None
        except json.JSONDecodeError as e:
            app_logger.warning(f"Redis数据格式错误（已删除）: {e}")
            # 删除损坏的数据
            try:
                self._redis.delete(self._generate_key(query, top_k, enable_self_rag, retrieval_policy))
            except:
                pass
            return None
        except Exception as e:
            app_logger.warning(f"Redis读取失败: {e}")
            return None

    def set(self, query: str, top_k: int, enable_self_rag: bool,
            result: Dict[str, Any], ttl: Optional[int] = None,
            retrieval_policy: Optional[Dict[str, Any]] = None):
        if not self._available or not self._redis:
            return

        try:
            key = self._generate_key(query, top_k, enable_self_rag, retrieval_policy)
            ttl = ttl or self.default_ttl
            # 使用JSON序列化替代pickle
            json_safe = self._json_safe_serialize(result)
            data = json.dumps(json_safe, ensure_ascii=False).encode('utf-8')
            self._redis.setex(key, ttl, data)
            app_logger.debug(f"Redis缓存设置: {key[:20]}... TTL={ttl}s")
        except Exception as e:
            app_logger.warning(f"Redis写入失败: {e}")

    def clear(self):
        if not self._available or not self._redis:
            return

        try:
            pattern = f"{self.key_prefix}*"
            keys = self._redis.keys(pattern)
            if keys:
                self._redis.delete(*keys)
                app_logger.info(f"Redis缓存已清除: {len(keys)}条")
        except Exception as e:
            app_logger.warning(f"Redis清除失败: {e}")

    def get_stats(self) -> Dict[str, Any]:
        if not self._available or not self._redis:
            return {"type": "unavailable"}

        try:
            info = self._redis.info("stats")
            pattern = f"{self.key_prefix}*"
            keys = self._redis.keys(pattern)
            return {
                "type": "redis",
                "keys_count": len(keys),
                "hits": info.get("keyspace_hits", 0),
                "misses": info.get("keyspace_misses", 0),
                "memory_used": info.get("used_memory_human", "unknown")
            }
        except Exception as e:
            return {"type": "error", "error": str(e)}


class Reranker:
    """文档重排序器，使用交叉编码器进行精细排序"""

    def __init__(self, model_name: str = "BAAI/bge-reranker-base"):
        self._model = None
        self._model_name = model_name
        self._available = False

    def _ensure_model(self):
        """延迟加载模型"""
        if self._model is None:
            try:
                self._model = CrossEncoder(self._model_name, max_length=512)
                self._available = True
                app_logger.info(f"Rerank模型加载成功: {self._model_name}")
            except Exception as e:
                app_logger.warning(f"Rerank模型加载失败: {e}")
                self._model = None
                self._available = False

    def rerank(self, query: str, documents: List[Dict[str, Any]],
               top_k: int = 5) -> List[Dict[str, Any]]:
        """
        对文档进行重排序

        Args:
            query: 查询文本
            documents: 初步检索的文档列表
            top_k: 返回前k个

        Returns:
            重排序后的文档列表
        """
        if not documents:
            return []

        self._ensure_model()

        if not self._available or not self._model:
            # 如果Rerank不可用，返回原始顺序
            return documents[:top_k]

        try:
            # 构建(query, doc)对
            doc_texts = [
                doc.get("content", "")
                for doc in documents
            ]
            pairs = [[query, doc_text] for doc_text in doc_texts]

            # 获取相关性分数
            scores = self._model.predict(pairs)

            # 按分数排序
            scored_docs = list(zip(documents, scores))
            scored_docs.sort(key=lambda x: x[1], reverse=True)

            # 返回排序后的结果
            reranked = []
            for doc, score in scored_docs[:top_k]:
                doc_copy = doc.copy()
                doc_copy["rerank_score"] = float(score)
                # 将L2距离转换为相似度（距离越小相似度越高）
                original_score = doc.get("similarity_score", 0)
                doc_copy["combined_score"] = (float(score) + (1 - min(original_score, 1))) / 2
                reranked.append(doc_copy)

            app_logger.info(f"Rerank完成: {len(documents)} -> {len(reranked)}个文档")
            return reranked

        except Exception as e:
            app_logger.warning(f"Rerank失败，回退到原始排序: {e}")
            return documents[:top_k]

    def is_available(self) -> bool:
        """检查模型是否可用"""
        self._ensure_model()
        return self._available


class BM25Retriever:
    """BM25关键词检索器（智能扩展版）"""
    
    def __init__(self, k1: float = 1.5, b: float = 0.75):
        """
        初始化BM25检索器
        
        Args:
            k1: 词频饱和参数，通常在1.2-2.0之间
            b: 文档长度归一化参数，通常为0.75
        """
        self.k1 = k1
        self.b = b
        self.documents = []  # 原始文档列表
        self.doc_texts = []  # 文档文本列表
        self.avg_doc_len = 0
        self.doc_freq = {}  # 词频统计
        self.total_docs = 0
        self.extracted_keywords = set()  # 自动提取的关键词
        self.use_jieba = False  # 是否使用jieba分词
    
    def _try_import_jieba(self):
        """尝试导入jieba分词"""
        if not self.use_jieba:
            try:
                import jieba
                self.jieba = jieba
                self.use_jieba = True
                app_logger.info("[BM25] 成功加载jieba分词")
            except ImportError:
                self.use_jieba = False
                app_logger.info("[BM25] 未安装jieba，使用简单分词")
    
    def _extract_keywords_from_documents(self):
        """从文档中自动提取关键词（增强版：特别处理产品型号）"""
        import re
        
        def simple_tokenize(text: str) -> List[str]:
            text = re.sub(r'[^\w\s]', ' ', text)
            tokens = []
            # 提取中文词
            chinese_tokens = re.findall(r'[\u4e00-\u9fff]+', text)
            for ct in chinese_tokens:
                # 简单按2-4字词提取
                for i in range(len(ct)):
                    for length in [2, 3, 4]:
                        if i + length <= len(ct):
                            tokens.append(ct[i:i+length])
            # 提取英文词（包括带数字的产品型号如 X12, X12Pro）
            english_tokens = re.findall(r'[a-zA-Z0-9]+', text.lower())
            tokens.extend(english_tokens)
            return [t for t in tokens if len(t) >= 2]
        
        all_keywords = set()
        for text in self.doc_texts:
            tokens = simple_tokenize(text)
            # 提取产品名模式（如：XX Pro, XX Plus, XX Max等）
            for token in tokens:
                # 跳过常见词
                if token in ['产品', '规格', '参数', '价格', '型号', '尺寸', '内存']:
                    continue
                all_keywords.add(token)
            
            # 特别提取产品型号模式（如 X12, X12 Pro, X12Pro 等）
            # 匹配：字母 + 数字 + 可选后缀 (Pro/Max/Plus/Air/Mini/SE)
            product_models = re.findall(
                r'[a-zA-Z]\d+(?:\s*(?:pro|max|plus|air|mini|se))?', 
                text.lower()
            )
            for model in product_models:
                model_lower = model.lower()
                all_keywords.add(model_lower)
                # 也添加无空格版本
                all_keywords.add(model_lower.replace(' ', ''))
                # 添加纯型号部分（如 x12pro -> x12）
                base_model = re.match(r'([a-zA-Z]\d+)', model_lower)
                if base_model:
                    all_keywords.add(base_model.group(1))
            
            # 提取"产品名称"行中的完整产品名
            product_names = re.findall(r'产品名称 [：:]\s*([^\n]+)', text)
            for name in product_names:
                name_lower = name.lower().strip()
                all_keywords.add(name_lower)
                # 提取其中的型号
                models_in_name = re.findall(r'[a-zA-Z]\d+(?:\s*(?:pro|max|plus|air|mini|se))?', name_lower)
                for model in models_in_name:
                    model_lower = model.lower()
                    all_keywords.add(model_lower)
                    all_keywords.add(model_lower.replace(' ', ''))
        
        self.extracted_keywords = all_keywords
        app_logger.info(f"[BM25] 自动提取了 {len(self.extracted_keywords)} 个关键词")
    
    # BM25停用词：产品无关的高频通用词（出现在几乎所有产品中）
    BM25_STOPWORDS = {
        # 通用产品属性词
        "支持", "产品介绍", "产品名称", "产品特点", "产品优势",
        "成本价", "标价", "最大优惠价", "序号",
        # 常见描述词
        "英寸", "小时", "续航", "蓝牙", "防水", "usb",
        "处理器", "内存", "存储", "屏幕", "电池",
        "高清", "无线", "智能", "专业", "标准",
        "版本", "型号", "重量", "尺寸", "颜色",
        # 价格相关（泛化太强）
        "价格", "多少钱", "多少", "元", "报价",
        # 其他高频无区别词
        "官网", "客服", "电话", "邮箱", "网址",
        "正品", "保障", "退货", "售后", "保修",
        "全新", "官方", "旗舰", "顶级",
    }

    def _expand_query(self, query: str) -> List[str]:
        """
        智能扩展查询词（自动提取关键词 + 产品型号特殊处理）

        Args:
            query: 原始查询文本

        Returns:
            扩展后的查询词列表
        """
        import re

        def tokenize(text: str) -> List[str]:
            text = re.sub(r'[^\w\s]', ' ', text.lower())
            return [w for w in text.split() if len(w) > 1]

        query_lower = query.lower().strip()
        tokens = tokenize(query)
        expanded_tokens = set(tokens)

        # 添加原始查询作为整体匹配（支持产品型号精确匹配如 x12, X12 Pro）
        expanded_tokens.add(query_lower)

        # 提取产品型号模式（如 X12, X12Pro, X12 Pro, x12等）
        # 匹配模式：字母+数字，可选后跟 Pro/Max/Plus/Air/Mini 等
        product_patterns = re.findall(r'[a-zA-Z]\d+(?:\s*(?:pro|max|plus|air|mini|se))?', query_lower)
        for pattern in product_patterns:
            expanded_tokens.add(pattern)
            # 也添加无空格版本（x12 pro -> x12pro）
            expanded_tokens.add(pattern.replace(' ', ''))

        # 1. 精确匹配：只在产品名称行中查找（高精准匹配）
        #    跳过模糊匹配，因为价格表产品名本身就是精确标识符
        #    保留产品型号扩展（如 x12 -> x12pro）

        # 2. 如果有jieba，使用更好的中文分词（仅用于原始查询词）
        if self.use_jieba:
            jieba_tokens = list(self.jieba.cut(query))
            expanded_tokens.update([t for t in jieba_tokens if len(t) >= 2])

        # 3. 过滤停用词：移除无区分度的通用词
        filtered = [t for t in expanded_tokens if t not in self.BM25_STOPWORDS]

        app_logger.info(f"[BM25] 查询扩展: '{query}' -> {filtered}")
        return filtered
        
    def index_documents(self, documents: List[Dict[str, Any]]):
        """
        为文档建立索引
        
        Args:
            documents: 文档列表，每个文档包含 'content' 字段
        """
        self.documents = documents
        self.doc_texts = [doc.get('content', '') for doc in documents]
        self.total_docs = len(self.documents)
        
        if self.total_docs == 0:
            return
        
        # 尝试加载jieba
        self._try_import_jieba()
        
        # 自动提取关键词
        self._extract_keywords_from_documents()
            
        # 分词（简单按空格和标点分词）
        def tokenize(text: str) -> List[str]:
            import re
            text = re.sub(r'[^\w\s]', ' ', text.lower())
            return [w for w in text.split() if len(w) > 1]
        
        # 计算文档长度
        doc_lens = [len(tokenize(text)) for text in self.doc_texts]
        self.avg_doc_len = sum(doc_lens) / self.total_docs if self.total_docs > 0 else 0
        
        # 统计词频
        self.doc_freq = {}
        for text in self.doc_texts:
            tokens = set(tokenize(text))
            for token in tokens:
                self.doc_freq[token] = self.doc_freq.get(token, 0) + 1
                
        app_logger.info(f"[BM25] 索引完成: {self.total_docs}个文档, {len(self.doc_freq)}个词条")
    
    def _calculate_idf(self, term: str) -> float:
        """计算IDF"""
        df = self.doc_freq.get(term, 0)
        if df == 0:
            return 0
        return math.log((self.total_docs - df + 0.5) / (df + 0.5) + 1)
    
    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        BM25检索（带同义词扩展）
        
        Args:
            query: 查询文本
            top_k: 返回前k个结果
            
        Returns:
            排序后的文档列表
        """
        if self.total_docs == 0 or not self.documents:
            return []
            
        def tokenize(text: str) -> List[str]:
            import re
            text = re.sub(r'[^\w\s]', ' ', text.lower())
            return [w for w in text.split() if len(w) > 1]
        
        # 使用同义词扩展查询
        query_tokens = self._expand_query(query)
        original_tokens = tokenize(query)
        
        # 确保所有查询词都是小写（与文档 tokenize 保持一致）
        query_tokens_lower = [token.lower() for token in query_tokens]
        
        app_logger.info(f"[BM25] 原查询词：{original_tokens} -> 扩展后：{query_tokens_lower}")
        
        if not query_tokens_lower:
            return []
            
        scores = []
        for i, text in enumerate(self.doc_texts):
            doc_tokens = tokenize(text)
            doc_len = len(doc_tokens)
            
            score = 0.0
            for term in query_tokens_lower:
                if term not in doc_tokens:
                    continue
                    
                # 计算词频
                tf = doc_tokens.count(term)
                
                # 计算IDF
                idf = self._calculate_idf(term)
                
                # BM25公式
                numerator = tf * (self.k1 + 1)
                denominator = tf + self.k1 * (1 - self.b + self.b * doc_len / max(self.avg_doc_len, 1))
                score += idf * (numerator / denominator)
                
            if score > 0:
                scores.append((score, i))
                
        # 按分数排序
        scores.sort(reverse=True)
        
        results = []
        for score, idx in scores[:top_k]:
            doc = self.documents[idx].copy()
            doc['bm25_score'] = score
            results.append(doc)
            
        return results
    
    def is_indexed(self) -> bool:
        """检查是否已建立索引"""
        return self.total_docs > 0


class LocalEmbeddings(Embeddings):
    """本地向量嵌入模型"""
    def __init__(self, model_name: str = None):
        self.model_name = model_name or settings.rag.embedding_model_name
        try:
            self.model = SentenceTransformer('shibing624/text2vec-base-chinese')
            app_logger.info(f"嵌入模型初始化完成: shibing624/text2vec-base-chinese (中文优化)")
        except Exception as e:
            app_logger.warning(f"加载中文模型失败，回退到英文模型: {e}")
            self.model = SentenceTransformer(self.model_name)
            app_logger.info(f"嵌入模型初始化完成: {self.model_name}")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """嵌入文档列表"""
        return self.model.encode(texts, normalize_embeddings=False).tolist()

    def embed_query(self, text: str) -> List[float]:
        """嵌入查询文本"""
        return self.model.encode([text], normalize_embeddings=False)[0].tolist()


class RAGQueryInput(BaseModel):
    """RAG检索输入参数"""
    query: str = Field(..., description="用户查询文本，用于向量检索")
    top_k: int = Field(default=3, description="返回最相似的文档数量")
    enable_self_rag: bool = Field(default=True, description="是否启用Self-RAG决策")


class CRAGDocumentInput(BaseModel):
    """CRAG文档入库输入参数"""
    file_path: str = Field(..., description="文档路径（PDF/Word/TXT）")
    metadata: Optional[Dict[str, Any]] = Field(default=None, description="文档元数据")


class RAGTool:
    """CRAG/Self-RAG工具类（使用 ChromaDB）"""
    name: str = "rag_retrieval"
    description: str = """
    基于ChromaDB向量数据库的检索工具，支持：
    1. Self-RAG：自主决策是否需要检索、检索多少文档；
    2. CRAG：持续更新知识库，动态检索最新文档；
    3. 混合检索：结合向量检索和关键词检索；
    4. 自适应检索：根据问题复杂度动态调整检索策略；
    5. Redis缓存：支持分布式Redis缓存；
    6. Rerank重排序：使用交叉编码器精细排序；
    7. 指标监控：实时监控系统运行状态。
    """
    args_schema: type[RAGQueryInput] = RAGQueryInput

    def __init__(self, use_redis: bool = True, use_rerank: bool = True):
        """初始化RAG工具"""
        self.embeddings = LocalEmbeddings()

        # 优化的分割器：优先按语义边界分割
        # 分隔符从粗粒度到细粒度排列
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.rag.chunk_size,
            chunk_overlap=settings.rag.chunk_overlap,
            # 优化分隔符顺序：段落 > 换行 > 中文句子 > 英文句子 > 其他
            separators=[
                # 1. 先按段落分割（两个换行）
                "\n\n",
                # 2. 中文段落（单个换行）
                "\n",
                # 3. 中文句子
                "。", "！", "？", "；",
                # 4. 英文句子
                ". ", "? ", "! ",
                # 5. 中文标点
                "，", "、",
                # 6. 英文标点和空格
                ", ", "; ", ": ", " ",
                # 7. 最后按字符分割
                ""
            ]
        )

        # ========== 三层架构初始化 ==========
        # Layer 1: 查询理解层 (使用 LLM 进行语义意图分类)
        # 注意：LLM 会在 retrieve 方法中动态传递，这里先用 None
        self.query_layer = QueryUnderstandingLayer(llm=None)
        # Layer 3: 生成上下文层
        self.generation_layer = GenerationContextLayer()
        # Layer 2: 检索层 (延迟初始化，详见下面的 @property)

        self.chroma_client = None
        self.collection = None
        self.vector_backend = None
        self._db_available = False
        self._collection_name = settings.vector_db.vector_db_collection_name

        # 刷新控制标记
        self._need_initial_load = True  # 标记是否需要初始加载
        self._refresh_interval = settings.rag.crag_refresh_interval  # 刷新间隔(秒)

        try:
            project_root = os.path.dirname(os.path.dirname(__file__))
            chroma_data_path = os.path.join(project_root, 'chroma_data')
            os.makedirs(chroma_data_path, exist_ok=True)

            self.chroma_client = chromadb.PersistentClient(
                path=chroma_data_path,
                settings=ChromaSettings(anonymized_telemetry=False)
            )

            # 尝试获取现有集合，如果存在就保留
            existing_count = 0
            try:
                self.collection = self.chroma_client.get_collection(self._collection_name)
                existing_count = self.collection.count()
                self._db_available = True
                app_logger.info(f"ChromaDB 使用已有集合: {self._collection_name}, 文档数: {existing_count}")
            except:
                # 集合不存在，创建新集合
                try:
                    self.collection = self.chroma_client.create_collection(
                        name=self._collection_name,
                        metadata={
                            "hnsw:space": "cosine",  # 使用余弦距离
                            "hnsw:construction_ef": 128,  # 构建时的搜索精度（越高越精确但越慢）
                            "hnsw:search_ef": 128,  # 搜索时的精度
                            "hnsw:M": 32  # 连接数（越高召回率越高但占内存越大）
                        }
                    )
                    self._db_available = True
                    app_logger.info(f"ChromaDB 创建新集合: {self._collection_name} (HNSW优化)")
                except Exception as e:
                    app_logger.warning(f"ChromaDB 创建集合失败: {e}")
                    self._db_available = False
                    self.collection = None
                    return

            # 如果集合为空，需要初始加载
            if existing_count == 0:
                self._need_initial_load = True
                app_logger.info("[INIT] 集合为空，将在首次检索时加载知识库")
            else:
                self._need_initial_load = False
                app_logger.info(f"[INIT] 集合已有 {existing_count} 个文档，跳过初始加载")

        except Exception as e:
            app_logger.warning(f"ChromaDB 初始化失败: {e}，RAG 功能将不可用")
            self.chroma_client = None
            self.collection = None
            self._db_available = False

        self.last_refresh_time = datetime.now() - timedelta(seconds=self._refresh_interval)

        # 初始化缓存
        self.metrics = RAGMetrics()
        if use_redis:
            # 优先使用Redis缓存
            redis_host = getattr(settings.redis, 'host', 'localhost')
            redis_port = getattr(settings.redis, 'port', 6379)
            redis_db = getattr(settings.redis, 'db', 0)
            redis_password = getattr(settings.redis, 'password', None)
            self.cache = RedisCache(
                host=redis_host,
                port=redis_port,
                db=redis_db,
                password=redis_password,
                key_prefix="rag:cache:",
                default_ttl=getattr(settings.rag, 'cache_ttl', 3600)
            )
            if not self.cache._available:
                app_logger.info("回退到内存缓存")
                self.cache = LocalCache(
                    max_size=getattr(settings.rag, 'cache_max_size', 10000),
                    default_ttl=getattr(settings.rag, 'cache_ttl', 86400)
                )
        else:
            self.cache = LocalCache(
                max_size=getattr(settings.rag, 'cache_max_size', 10000),
                default_ttl=getattr(settings.rag, 'cache_ttl', 86400)
            )

        # 初始化Reranker
        self.reranker = Reranker() if use_rerank else None
        
        # 初始化专业BM25检索器（使用ProfessionalBM25）
        self.bm25_retriever = ProfessionalBM25(k1=1.5, b=0.75, use_jieba=True)
        self._bm25_indexed = False
        self._bm25_doc_count = 0  # 缓存的文档数量

        # ========== 三层架构初始化 (Layer 2: 检索层) ==========
        # 延迟初始化，因为在初始化时 ChromaDB 可能还未准备好
        self._retrieval_layer = None

        self.executor = ThreadPoolExecutor(max_workers=4)

        app_logger.info("CRAG/Self-RAG工具（ChromaDB版 + 三层架构）初始化完成")

    @property
    def retrieval_layer(self):
        """延迟加载检索层"""
        if self._retrieval_layer is None:
            from tools.rag.retrieval_context import RetrievalLayer
            self._retrieval_layer = RetrievalLayer(self)
        return self._retrieval_layer

    def _get_cache_key(
        self,
        query: str,
        top_k: int,
        enable_self_rag: bool,
        retrieval_policy: Optional[Dict[str, Any]] = None,
    ) -> str:
        return build_retrieval_cache_key(
            query,
            top_k,
            enable_self_rag,
            retrieval_policy=retrieval_policy,
        )

    def _detect_content_type(self, file_path: str) -> str:
        """
        根据文件名检测内容类型

        Returns:
            "faq", "price_list", "manual", 或 "general"
        """
        filename = os.path.basename(file_path).lower()
        if "faq" in filename or "q&a" in filename or "常见问题" in filename:
            return "faq"
        elif "价格" in filename or "price" in filename:
            return "price_list"
        elif "手册" in filename or "manual" in filename or "指南" in filename:
            return "manual"
        return "general"

    # 内容类型感知分块策略
    CHUNKING_STRATEGIES = {
        "faq": {"chunk_size": 180, "chunk_overlap": 35},         # 短问答对
        "price_list": {"chunk_size": 250, "chunk_overlap": 30},  # 价格条目：增大确保完整产品为1个chunk
        "manual": {"chunk_size": 450, "chunk_overlap": 90},     # 详细手册
        "general": {"chunk_size": 400, "chunk_overlap": 50}     # 默认
    }

    def _detect_document_strategy(self, file_path: str) -> str:
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext in {".html", ".htm", ".docx"}:
            return "structured_document"
        if file_ext == ".xlsx":
            return "table_document"
        if file_ext == ".pdf" and self._pdf_has_outline(file_path):
            return "structured_document"
        return "recursive_text"

    def _pdf_has_outline(self, file_path: str) -> bool:
        try:
            with pdfplumber.open(file_path) as pdf:
                catalog = getattr(getattr(pdf, "doc", None), "catalog", None)
                if hasattr(catalog, "get"):
                    return bool(catalog.get("Outlines"))
        except Exception:
            return False
        return False

    def _build_recursive_splitter(self, chunk_size: int, chunk_overlap: int,
                                  content_type: str) -> RecursiveCharacterTextSplitter:
        if content_type == "price_list":
            separators = [
                "\n===== ",
                "\n\n",
                "\n",
                "。",
                "？",
                "！",
                ". ",
                "? ",
                "! ",
                "，",
                "；",
                "：",
                ", ",
                "; ",
                ": ",
                " ",
                "",
            ]
        else:
            separators = [
                "\n\n",
                "\n",
                "。",
                "？",
                "！",
                ". ",
                "? ",
                "! ",
                "，",
                "；",
                "：",
                ", ",
                "; ",
                ": ",
                " ",
                "",
            ]

        return RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators
        )

    def _read_recursive_text(self, file_path: str) -> str:
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        if file_ext == ".docx":
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs if para.text])
        if file_ext in {".txt", ".html", ".htm"}:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        if file_ext == ".xlsx":
            rows = []
            for sheet_name, sheet_rows in self._load_xlsx_rows(file_path):
                rows.append(f"[Sheet] {sheet_name}")
                for row in sheet_rows:
                    values = [
                        str(cell).strip()
                        for cell in row
                        if cell is not None and str(cell).strip() != ""
                    ]
                    if values:
                        rows.append("\t".join(values))
            return "\n".join(rows)

        raise ValueError(f"涓嶆敮鎸佺殑鏂囦欢绫诲瀷: {file_path}")

    def _load_xlsx_rows(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return self._load_xlsx_rows_from_archive(file_path)

        workbook = load_workbook(file_path, read_only=True, data_only=True)
        try:
            return [
                (sheet.title, [list(row) for row in sheet.iter_rows(values_only=True)])
                for sheet in workbook.worksheets
            ]
        finally:
            workbook.close()

    def _load_xlsx_rows_from_archive(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        with zipfile.ZipFile(file_path) as archive:
            workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
            relationships_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
            relationship_targets: Dict[str, str] = {}

            for relationship in relationships_root.findall(f"{{{_XLSX_REL_NS}}}Relationship"):
                relationship_id = relationship.attrib.get("Id")
                target = relationship.attrib.get("Target")
                if not relationship_id or not target:
                    continue
                relationship_targets[relationship_id] = posixpath.normpath(
                    posixpath.join("xl", target)
                ).lstrip("/")

            shared_strings = self._load_xlsx_shared_strings(archive)
            sheets: List[Tuple[str, List[List[Any]]]] = []
            for sheet in workbook_root.findall(f".//{{{_XLSX_MAIN_NS}}}sheet"):
                sheet_name = sheet.attrib.get("name") or "Sheet"
                relationship_id = sheet.attrib.get(f"{{{_XLSX_DOC_REL_NS}}}id")
                sheet_path = relationship_targets.get(relationship_id or "")
                if not sheet_path:
                    continue
                sheets.append(
                    (
                        sheet_name,
                        self._parse_xlsx_sheet_rows(archive.read(sheet_path), shared_strings),
                    )
                )
            return sheets

    def _load_xlsx_shared_strings(self, archive: zipfile.ZipFile) -> List[str]:
        if "xl/sharedStrings.xml" not in archive.namelist():
            return []

        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
        shared_strings: List[str] = []
        for string_item in root.findall(f".//{{{_XLSX_MAIN_NS}}}si"):
            text_parts = [
                text_node.text or ""
                for text_node in string_item.findall(f".//{{{_XLSX_MAIN_NS}}}t")
            ]
            shared_strings.append("".join(text_parts))
        return shared_strings

    def _parse_xlsx_sheet_rows(self, sheet_xml: bytes, shared_strings: List[str]) -> List[List[Any]]:
        root = ET.fromstring(sheet_xml)
        rows: List[List[Any]] = []
        max_columns = 0

        for row_node in root.findall(f".//{{{_XLSX_MAIN_NS}}}sheetData/{{{_XLSX_MAIN_NS}}}row"):
            row_values: Dict[int, Any] = {}
            row_max_column = 0

            for cell_index, cell_node in enumerate(
                row_node.findall(f"{{{_XLSX_MAIN_NS}}}c"),
                start=1,
            ):
                column_index = self._xlsx_column_index(cell_node.attrib.get("r", "")) or cell_index
                row_values[column_index] = self._parse_xlsx_cell_value(cell_node, shared_strings)
                row_max_column = max(row_max_column, column_index)

            if row_max_column == 0:
                rows.append([])
                continue

            max_columns = max(max_columns, row_max_column)
            rows.append([row_values.get(index) for index in range(1, row_max_column + 1)])

        if max_columns > 0:
            for row in rows:
                if len(row) < max_columns:
                    row.extend([None] * (max_columns - len(row)))

        return rows

    def _xlsx_column_index(self, cell_reference: str) -> int:
        match = re.match(r"([A-Z]+)", (cell_reference or "").upper())
        if not match:
            return 0

        column_index = 0
        for letter in match.group(1):
            column_index = (column_index * 26) + (ord(letter) - ord("A") + 1)
        return column_index

    def _parse_xlsx_cell_value(self, cell_node: ET.Element, shared_strings: List[str]) -> Any:
        cell_type = cell_node.attrib.get("t")

        if cell_type == "inlineStr":
            text_parts = [
                text_node.text or ""
                for text_node in cell_node.findall(f".//{{{_XLSX_MAIN_NS}}}t")
            ]
            return "".join(text_parts)

        value_node = cell_node.find(f"{{{_XLSX_MAIN_NS}}}v")
        if value_node is None or value_node.text is None:
            return None

        raw_value = value_node.text
        if cell_type == "s":
            try:
                return shared_strings[int(raw_value)]
            except (ValueError, IndexError):
                return raw_value
        if cell_type == "b":
            return "TRUE" if raw_value == "1" else "FALSE"

        try:
            numeric_value = float(raw_value)
        except ValueError:
            return raw_value

        if numeric_value.is_integer():
            return int(numeric_value)
        return numeric_value

    def _load_recursive_document(self, file_path: str, chunk_size: int, chunk_overlap: int,
                                 metadata: Dict[str, Any]) -> List[Document]:
        text = self._read_recursive_text(file_path)
        temp_splitter = self._build_recursive_splitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            content_type=metadata.get("content_type", "general")
        )
        chunks = temp_splitter.split_text(text)
        return [Document(page_content=chunk, metadata=dict(metadata)) for chunk in chunks]

    def _extract_html_blocks(self, file_path: str) -> List[Dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8") as f:
            html_text = f.read()

        parser = _StructuredHTMLBlockParser()
        parser.feed(html_text)
        parser.close()
        return parser.blocks

    def _extract_docx_blocks(self, file_path: str) -> List[Dict[str, Any]]:
        document = docx.Document(file_path)
        blocks: List[Dict[str, Any]] = []
        heading_stack: List[str] = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue

            style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
            heading_match = re.match(r"Heading\s+(\d+)", style_name, re.IGNORECASE)

            if heading_match:
                level = max(1, int(heading_match.group(1)))
                heading_stack = heading_stack[: level - 1]
                heading_stack.append(paragraph_text)
                continue

            blocks.append(
                {
                    "heading_path": list(heading_stack),
                    "text": paragraph_text,
                }
            )

        return blocks

    def _chunk_structured_blocks(self, blocks: List[Dict[str, Any]], chunk_size: int,
                                 chunk_overlap: int, metadata: Dict[str, Any]) -> List[Document]:
        documents: List[Document] = []
        content_type = metadata.get("content_type", "general")

        for block_index, block in enumerate(blocks):
            block_text = (block.get("text") or "").strip()
            if not block_text:
                continue

            heading_path = [part.strip() for part in block.get("heading_path", []) if part and part.strip()]
            heading_label = " > ".join(heading_path)

            if heading_label:
                prefix = f"{heading_label}\n"
                body_chunk_size = max(1, chunk_size - len(prefix))
            else:
                prefix = ""
                body_chunk_size = chunk_size

            body_chunk_overlap = min(chunk_overlap, max(body_chunk_size - 1, 0))
            splitter = self._build_recursive_splitter(
                chunk_size=body_chunk_size,
                chunk_overlap=body_chunk_overlap,
                content_type=content_type
            )
            body_chunks = splitter.split_text(block_text) or [block_text]

            for chunk_part_index, body_chunk in enumerate(body_chunks):
                chunk_metadata = dict(metadata)
                if heading_label:
                    chunk_metadata["heading_path"] = heading_label
                chunk_metadata["structured_block_index"] = block_index
                chunk_metadata["structured_chunk_index"] = chunk_part_index
                documents.append(
                    Document(
                        page_content=f"{prefix}{body_chunk}" if prefix else body_chunk,
                        metadata=chunk_metadata,
                    )
                )

        return documents

    def _load_structured_document(self, file_path: str, chunk_size: int, chunk_overlap: int,
                                  metadata: Dict[str, Any]) -> List[Document]:
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext in {".html", ".htm"}:
            blocks = self._extract_html_blocks(file_path)
            return self._chunk_structured_blocks(blocks, chunk_size, chunk_overlap, metadata)
        if file_ext == ".docx":
            blocks = self._extract_docx_blocks(file_path)
            return self._chunk_structured_blocks(blocks, chunk_size, chunk_overlap, metadata)

        return self._load_recursive_document(file_path, chunk_size, chunk_overlap, metadata)

    def _load_table_document(self, file_path: str, chunk_size: int, chunk_overlap: int,
                             metadata: Dict[str, Any]) -> List[Document]:
        documents: List[Document] = []
        overlap_rows = 1 if chunk_overlap > 0 else 0

        for sheet_name, sheet_rows in self._load_xlsx_rows(file_path):
            if not sheet_rows:
                continue

            max_columns = max((len(row) for row in sheet_rows), default=0)
            if max_columns == 0:
                continue

            header_row = list(sheet_rows[0]) + [None] * (max_columns - len(sheet_rows[0]))
            headers = [
                str(value).strip() if value is not None and str(value).strip() != "" else f"Column {index + 1}"
                for index, value in enumerate(header_row)
            ]
            body_rows = [
                (
                    row_number,
                    list(row) + [None] * (max_columns - len(row)),
                )
                for row_number, row in enumerate(sheet_rows[1:], start=2)
                if any(cell is not None and str(cell).strip() != "" for cell in row)
            ]
            if not body_rows:
                continue

            start_index = 0
            while start_index < len(body_rows):
                chunk_row_lines: List[str] = []
                end_index = start_index

                while end_index < len(body_rows):
                    row_number, row_values = body_rows[end_index]
                    value_text = " | ".join(
                        "" if cell is None else str(cell).strip()
                        for cell in row_values
                    )
                    candidate_lines = chunk_row_lines + [f"{row_number}: {value_text}"]
                    candidate_content = "\n".join(
                        [
                            f"Sheet: {sheet_name}",
                            f"Headers: {' | '.join(headers)}",
                            f"Rows {body_rows[start_index][0]}-{row_number}",
                            *candidate_lines,
                        ]
                    )
                    if chunk_row_lines and len(candidate_content) > chunk_size:
                        break
                    chunk_row_lines = candidate_lines
                    end_index += 1

                if not chunk_row_lines:
                    row_number, row_values = body_rows[start_index]
                    chunk_row_lines = [
                        f"{row_number}: "
                        + " | ".join("" if cell is None else str(cell).strip() for cell in row_values)
                    ]
                    end_index = start_index + 1

                start_row = body_rows[start_index][0]
                end_row = body_rows[end_index - 1][0]
                chunk_metadata = dict(metadata)
                chunk_metadata["sheet_name"] = sheet_name
                chunk_metadata["row_range"] = f"{start_row}-{end_row}"
                documents.append(
                    Document(
                        page_content="\n".join(
                            [
                                f"Sheet: {sheet_name}",
                                f"Headers: {' | '.join(headers)}",
                                f"Rows {start_row}-{end_row}",
                                *chunk_row_lines,
                            ]
                        ),
                        metadata=chunk_metadata,
                    )
                )

                if end_index >= len(body_rows):
                    break
                if overlap_rows:
                    start_index = max(start_index + 1, end_index - overlap_rows)
                else:
                    start_index = end_index

        return documents

    def load_document(self, file_path: str, chunk_size: int = None, chunk_overlap: int = None,
                       metadata: Optional[Dict[str, Any]] = None,
                       content_type_hint: str = None) -> List[Document]:
        """加载本地文档（PDF/Word/TXT）

        Args:
            file_path: 文档路径
            chunk_size: 文档块大小（可选，默认使用配置值）
            chunk_overlap: 块重叠大小（可选，默认使用配置值）
            metadata: 额外的元数据
            content_type_hint: 内容类型提示 ("faq", "price_list", "manual", "general")
        """
        if not os.path.exists(file_path):
            raise ValueError(f"文档不存在: {file_path}")

        metadata = dict(metadata or {})
        metadata["file_path"] = file_path
        metadata["update_time"] = datetime.now().isoformat()

        # 内容类型检测
        content_type = content_type_hint or self._detect_content_type(file_path)
        metadata["content_type"] = content_type

        # 如果没有提供参数，使用内容类型策略
        if chunk_size is None:
            chunk_size = self.text_splitter._chunk_size
        if chunk_overlap is None:
            chunk_overlap = self.text_splitter._chunk_overlap
        document_strategy = self._detect_document_strategy(file_path)
        metadata["document_strategy"] = document_strategy

        if document_strategy == "structured_document":
            documents = self._load_structured_document(
                file_path=file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                metadata=metadata
            )
        elif document_strategy == "table_document":
            documents = self._load_table_document(
                file_path=file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                metadata=metadata
            )
        else:
            documents = self._load_recursive_document(
                file_path=file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                metadata=metadata
            )

        app_logger.info(
            f"鏂囨。鍔犺浇瀹屾垚: {file_path}锛屽垏鍒嗕负{len(documents)}鍧?"
            f" (type={content_type}, strategy={document_strategy}, size={chunk_size}, overlap={chunk_overlap})"
        )
        self.metrics.record_documents_processed(len(documents))
        return documents

    def _ensure_collection_exists(self):
        """确保集合存在，如果不存在则创建"""
        if not self._db_available or self.chroma_client is None:
            raise RuntimeError("向量数据库客户端不可用")
        
        try:
            # 尝试获取集合
            self.collection = self.chroma_client.get_collection(self._collection_name)
            app_logger.debug(f"集合 {self._collection_name} 已存在")
        except Exception:
            # 集合不存在，创建新集合
            try:
                self.collection = self.chroma_client.create_collection(
                    name=self._collection_name,
                    metadata={"hnsw:space": "cosine"}
                )
                app_logger.info(f"集合 {self._collection_name} 创建成功")
            except Exception as e:
                app_logger.error(f"创建集合失败: {e}")
                raise RuntimeError(f"无法创建集合: {e}")
    
    def _refresh_vector_backend(self) -> None:
        if self._db_available and self.collection is not None:
            self.vector_backend = ChromaVectorStoreBackend(
                collection=self.collection,
                embeddings=self.embeddings,
                available=True,
            )
        else:
            self.vector_backend = None

    def add_documents_to_vector_db(self, documents: List[Document]) -> List[str]:
        """将文档写入向量数据库（含去重机制）

        增强版：在写入前进行content_hash去重检查，避免重复导入导致文档数量异常增多
        """
        import hashlib

        # 确保集合存在
        self._ensure_collection_exists()

        if not self._db_available or self.collection is None:
            raise RuntimeError("向量数据库不可用，无法添加文档")

        texts = [doc.page_content for doc in documents]
        embeddings = self.embeddings.embed_documents(texts)

        # 使用时间戳+随机数生成唯一ID
        import uuid
        import time
        base_id = f"doc_{int(time.time() * 1000)}"

        # 获取源文件路径
        source_file = None
        for doc in documents:
            if doc.metadata and doc.metadata.get("file_path"):
                source_file = doc.metadata["file_path"]
                break

        # ========== 去重机制：计算content_hash并检查重复 ==========
        app_logger.info(f"[去重] 开始对 {len(documents)} 个文档块执行去重检查...")

        try:
            # 获取数据库中现有文档的元数据（用于去重比对）
            existing_docs = self.collection.get(include=["metadatas", "documents"])
            existing_hashes = set()

            # 构建已存在文档的hash集合
            for i, metadata in enumerate(existing_docs.get("metadatas", [])):
                existing_hash = metadata.get("content_hash")
                if existing_hash:
                    existing_hashes.add(existing_hash)
                    # 同时记录该hash对应的源文件
                    existing_source = metadata.get("source_file", "")

            app_logger.info(f"[去重] 数据库中已有 {len(existing_hashes)} 个唯一content_hash")

            # 过滤掉重复文档
            unique_documents = []
            duplicate_count = 0
            skipped_hashes = set()

            for doc, embedding in zip(documents, embeddings):
                # 生成当前文档块的content_hash（基于内容+源文件路径）
                content_for_hash = f"{doc.page_content}|{source_file or 'unknown'}"
                content_hash = hashlib.md5(content_for_hash.encode('utf-8')).hexdigest()[:16]

                if content_hash in existing_hashes or content_hash in skipped_hashes:
                    duplicate_count += 1
                    if content_hash not in skipped_hashes:
                        app_logger.warning(f"[去重] 发现重复文档块 (hash: {content_hash})，已跳过")
                        skipped_hashes.add(content_hash)
                    continue

                # 标记此hash为已处理
                skipped_hashes.add(content_hash)
                unique_documents.append((doc, embedding, content_hash))

            if duplicate_count > 0:
                app_logger.warning(f"[去重] ✓ 共过滤 {duplicate_count} 个重复文档块，实际写入 {len(unique_documents)} 个")
            else:
                app_logger.info(f"[去重] ✓ 无重复文档，全部 {len(unique_documents)} 个文档块将写入")

            # 如果全部重复，直接返回
            if not unique_documents:
                app_logger.warning("[去重] 所有文档块均为重复，跳过写入")
                return []

        except Exception as dedup_err:
            # 去重检查失败时，回退到原始逻辑（不阻止写入）
            app_logger.error(f"[去重] 去重检查失败，回退到原始逻辑: {dedup_err}")
            unique_documents = [(doc, embedding, hashlib.md5(f"{doc.page_content}|{source_file or 'unknown'}".encode('utf-8')).hexdigest()[:16])
                               for doc, embedding in zip(documents, embeddings)]

        # ========== 执行文档写入 ==========
        ids = []
        for i, (doc, embedding, content_hash) in enumerate(unique_documents):
            doc_id = f"{base_id}_{uuid.uuid4().hex[:8]}"
            ids.append(doc_id)

            # 为每个文档块添加完整的元数据（包含content_hash用于后续去重）
            metadata = doc.metadata.copy() if doc.metadata else {}
            metadata["chunk_id"] = doc_id
            metadata["chunk_index"] = i
            metadata["source_file"] = source_file or metadata.get("file_path", "未知")
            metadata["content_hash"] = content_hash  # 新增：存储content_hash供后续去重使用

            self.collection.add(
                ids=[doc_id],
                embeddings=[embedding],
                documents=[doc.page_content],
                metadatas=[metadata]
            )

        app_logger.info(f"{len(unique_documents)}个文档块写入向量数据库（已去重），来源: {source_file}")
        self._invalidate_cache()
        return ids

    def _invalidate_cache(self):
        """清除缓存"""
        self.cache.clear()
        self._bm25_documents = []
        self._bm25_indexed = False  # 标记需要重建BM25索引
        app_logger.debug("缓存已清除（文档已更新）")

    def delete_documents_by_source(self, source_file: str) -> int:
        """删除指定源文件的所有文档块

        Args:
            source_file: 源文件路径

        Returns:
            删除的文档数量
        """
        if not self._db_available or self.collection is None:
            raise RuntimeError("向量数据库不可用，无法删除文档")

        try:
            # 获取所有文档
            all_docs = self.collection.get(include=["metadatas"])
            ids_to_delete = []

            for i, metadata in enumerate(all_docs.get("metadatas", [])):
                if metadata.get("source_file") == source_file:
                    ids_to_delete.append(all_docs["ids"][i])

            if ids_to_delete:
                self.collection.delete(ids=ids_to_delete)
                app_logger.info(f"删除源文件 '{source_file}' 的 {len(ids_to_delete)} 个文档块")

                # 关键修复：清除缓存并强制标记BM25索引需要重建
                self._invalidate_cache()
                self._bm25_indexed = False  # 强制下次检索时重建BM25索引
                self._bm25_doc_count = 0   # 重置缓存的文档数量
                app_logger.info(f"[删除] 已标记BM25索引需重建（源文件已删除: {source_file}）")

                return len(ids_to_delete)
            else:
                app_logger.warning(f"未找到源文件 '{source_file}' 对应的文档")
                return 0

        except Exception as e:
            app_logger.error(f"删除文档失败: {str(e)}")
            return 0

    def delete_document(self, document_id: str) -> bool:
        """删除指定文档"""
        if not self._db_available or self.collection is None:
            raise RuntimeError("向量数据库不可用，无法删除文档")
        try:
            self.collection.delete(ids=[document_id])
            app_logger.info(f"文档删除成功: {document_id}")

            # 关键修复：清除缓存并强制标记BM25索引需要重建
            self._invalidate_cache()
            self._bm25_indexed = False  # 强制下次检索时重建BM25索引
            self._bm25_doc_count = 0   # 重置缓存的文档数量

            return True
        except Exception as e:
            app_logger.error(f"文档删除失败: {str(e)}")
            return False

    def update_document(self, document_id: str, content: str, metadata: Dict[str, Any]) -> bool:
        """更新指定文档"""
        if not self._db_available or self.collection is None:
            raise RuntimeError("向量数据库不可用，无法更新文档")
        try:
            self.collection.delete(ids=[document_id])
            embedding = self.embeddings.embed_query(content)
            self.collection.add(
                ids=[document_id],
                embeddings=[embedding],
                documents=[content],
                metadatas=[metadata]
            )
            app_logger.info(f"文档更新成功: {document_id}")
            self._invalidate_cache()
            return True
        except Exception as e:
            app_logger.error(f"文档更新失败: {str(e)}")
            return False

    def force_refresh(self) -> None:
        """强制刷新知识库（用于手动更新文档后）"""
        app_logger.info("[force_refresh] 强制刷新知识库...")
        self._need_initial_load = True
        self.last_refresh_time = datetime.now() - timedelta(seconds=self._refresh_interval)
        self.refresh_crag_knowledge_base()

    def clear_and_rebuild_collection(self) -> bool:
        """彻底清空并重建集合（用于知识库重载）

        通过 chroma_client 直接删除并重建，确保旧数据完全清除
        """
        if self.chroma_client is None:
            app_logger.error("[clear_and_rebuild] ChromaDB 客户端不可用")
            return False

        try:
            # 通过 client 直接删除集合（最可靠）
            self.chroma_client.delete_collection(self._collection_name)
            app_logger.info(f"[clear_and_rebuild] 集合 {self._collection_name} 已删除")

            # 重建空集合
            self.collection = self.chroma_client.create_collection(
                name=self._collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            self._db_available = True
            self._invalidate_cache()
            app_logger.info(f"[clear_and_rebuild] 集合已重建为空集合，文档数: {self.collection.count()}")
            return True
        except Exception as e:
            app_logger.error(f"[clear_and_rebuild] 清空重建集合失败: {e}")
            self._db_available = False
            self.collection = None
            return False

    def ensure_collection_ready(self, max_retries: int = 3) -> bool:
        """确保集合可用，如果不存在则重建并加载数据

        这个方法在每次检索失败后可以调用来恢复集合状态。
        增强版：支持重试机制、完整的HNSW参数配置、详细的错误诊断

        Args:
            max_retries: 最大重试次数，默认3次
        """
        if self.chroma_client is None:
            app_logger.error("[ensure_collection_ready] ChromaDB 客户端不可用")
            return False

        for attempt in range(max_retries):
            try:
                app_logger.info(f"[ensure_collection_ready] 第{attempt + 1}次尝试检查集合状态...")

                # 步骤1：检查集合是否存在
                try:
                    self.collection = self.chroma_client.get_collection(self._collection_name)
                    count = self.collection.count()
                    self._db_available = True
                    app_logger.info(f"[ensure_collection_ready] ✓ 集合存在且正常，名称: {self._collection_name}, 文档数: {count}")
                    return True
                except Exception as get_err:
                    app_logger.warning(f"[ensure_collection_ready] 集合获取失败 (尝试 {attempt + 1}/{max_retries}): {type(get_err).__name__}: {get_err}")

                # 步骤2：集合不存在或损坏，执行清理和重建
                app_logger.warning(f"[ensure_collection_ready] 执行集合清理和重建流程...")

                # 尝试删除可能存在的损坏集合
                try:
                    existing_collections = self.chroma_client.list_collections()
                    collection_names = [c.name for c in existing_collections]
                    app_logger.info(f"[ensure_collection_ready] 当前已有集合: {collection_names}")

                    if self._collection_name in collection_names:
                        app_logger.warning(f"[ensure_collection_ready] 发现同名但无法访问的集合，强制删除...")
                        self.chroma_client.delete_collection(self._collection_name)
                        import time
                        time.sleep(0.5)  # 等待文件系统同步
                except Exception as del_err:
                    app_logger.warning(f"[ensure_collection_ready] 清理旧集合时出错（可忽略）: {del_err}")

                # 步骤3：创建新集合（使用与初始化相同的完整HNSW参数）
                self.collection = self.chroma_client.create_collection(
                    name=self._collection_name,
                    metadata={
                        "hnsw:space": "cosine",
                        "hnsw:construction_ef": 128,
                        "hnsw:search_ef": 128,
                        "hnsw:M": 32,
                        "created_at": datetime.now().isoformat(),
                        "rebuild_reason": "auto_recovery"
                    }
                )
                self._db_available = True
                app_logger.info(f"[ensure_collection_ready] ✓ 新集合已创建: {self._collection_name} (含完整HNSW配置)")

                # 步骤4：清除所有缓存并标记需要重建索引
                self._invalidate_cache()
                self._bm25_indexed = False
                self._bm25_doc_count = 0

                # 步骤5：从磁盘重新加载文档到集合
                app_logger.info("[ensure_collection_ready] 开始从磁盘加载知识库文档...")
                load_result = self.refresh_crag_knowledge_base()
                final_count = self.collection.count() if self.collection else 0
                app_logger.info(f"[ensure_collection_ready] ✓ 集合恢复完成！最终文档数: {final_count}")
                return True

            except Exception as e:
                app_logger.error(f"[ensure_collection_ready] 第{attempt + 1}次尝试失败: {type(e).__name__}: {e}")
                if attempt < max_retries - 1:
                    import time
                    time.sleep(1 * (attempt + 1))  # 递增等待时间
                continue

        app_logger.error(f"[ensure_collection_ready] ✗ 经过{max_retries}次尝试仍无法恢复集合")
        self._db_available = False
        self.collection = None
        return False

    def refresh_crag_knowledge_base(self) -> None:
        """CRAG：定时刷新知识库"""
        now = datetime.now()
        if now - self.last_refresh_time < timedelta(seconds=settings.rag.crag_refresh_interval):
            app_logger.debug("CRAG知识库未到刷新时间，跳过")
            return

        if not self._db_available:
            app_logger.warning("向量数据库不可用，跳过知识库刷新")
            return
        
        # 确保集合存在
        try:
            self._ensure_collection_exists()
            app_logger.info("集合检查完成，开始刷新CRAG知识库...")
        except Exception as e:
            app_logger.error(f"集合检查失败，跳过知识库刷新: {e}")
            return
        data_sources = settings.rag.get_absolute_data_paths()

        total_docs = 0
        for source in data_sources:
            app_logger.info(f"检查数据源: {source}")
            if os.path.exists(source):
                if os.path.isdir(source):
                    file_count = 0
                    for file in os.listdir(source):
                        if file.lower().endswith((".pdf", ".docx", ".txt")):
                            file_path = os.path.join(source, file)
                            try:
                                docs = self.load_document(file_path)
                                self.add_documents_to_vector_db(docs)
                                total_docs += len(docs)
                                file_count += 1
                            except Exception as e:
                                app_logger.error(f"处理文件 {file} 时出错: {str(e)}")
                    app_logger.info(f"目录 {source} 处理完成，共处理 {file_count} 个文件")
                else:
                    try:
                        docs = self.load_document(source)
                        self.add_documents_to_vector_db(docs)
                        total_docs += len(docs)
                    except Exception as e:
                        app_logger.error(f"处理文件 {source} 时出错: {str(e)}")
            elif source.startswith("http"):
                app_logger.warning("URL数据源暂未实现")
            else:
                app_logger.warning(f"数据源不存在或无法访问: {source}")

        self.last_refresh_time = now
        app_logger.info(f"CRAG知识库刷新完成，共导入 {total_docs} 个文档块")

    def _analyze_query_complexity(self, query: str, llm) -> Dict[str, Any]:
        """分析查询复杂度"""
        complexity_prompt = f"""
        分析以下问题的复杂度：
        问题：{query}

        判断标准：
        - 简单问题：事实性问题，答案明确
        - 中等问题：需要综合多个信息点
        - 复杂问题：需要深入分析、多步骤推理

        返回JSON格式（仅返回JSON）：
        {{
            "complexity": "simple/medium/complex",
            "estimated_docs_needed": 1-10,
            "reason": "判断理由"
        }}
        """

        try:
            response = llm.invoke(complexity_prompt)
            analysis = json.loads(response.content.strip())
            return analysis
        except Exception as e:
            return {
                "complexity": "medium",
                "estimated_docs_needed": 3,
                "reason": f"分析失败: {str(e)}"
            }

    def _get_self_rag_cache_key(self, query: str) -> str:
        """生成Self-RAG决策缓存Key"""
        return hashlib.md5(query.encode('utf-8')).hexdigest()

    def self_rag_decision(self, query: str, llm, use_cache: bool = True) -> Dict[str, Any]:
        """Self-RAG决策逻辑（带缓存优化）"""
        # 检查缓存
        if use_cache and hasattr(self, '_self_rag_cache'):
            cache_key = self._get_self_rag_cache_key(query)
            if cache_key in self._self_rag_cache:
                cached = self._self_rag_cache[cache_key]
                cached_time = cached.get('_cached_at', 0)
                if time.time() - cached_time < 3600:  # 缓存1小时
                    app_logger.debug(f"[Self-RAG] 缓存命中: '{query}'")
                    return cached
        
        # 初始化缓存
        if not hasattr(self, '_self_rag_cache'):
            self._self_rag_cache = {}
        
        complexity_analysis = self._analyze_query_complexity(query, llm)

        complexity = complexity_analysis.get("complexity", "medium")
        estimated_docs = complexity_analysis.get("estimated_docs_needed", 3)

        if complexity == "simple":
            default_top_k = 1
            default_threshold = 0.6
        elif complexity == "complex":
            default_top_k = 5
            default_threshold = 0.8
        else:
            default_top_k = 3
            default_threshold = 0.7

        decision_prompt = f"""
        判断是否需要通过向量检索来回答以下问题：
        问题：{query}

        决策规则：
        1. 如果问题可以通过内置知识准确回答，无需检索；
        2. 如果问题涉及特定文档、最新数据、产品细节，需要检索；

        返回JSON格式：
        {{
            "need_retrieval": true/false,
            "confidence": 0.0-1.0,
            "reason": "决策原因"
        }}
        """

        try:
            response = llm.invoke(decision_prompt)
            decision = json.loads(response.content.strip())
            decision["complexity"] = complexity
            decision["adaptive_top_k"] = estimated_docs
            decision["_cached_at"] = time.time()
            
            # 缓存结果
            cache_key = self._get_self_rag_cache_key(query)
            self._self_rag_cache[cache_key] = decision
            
            # 限制缓存大小
            if len(self._self_rag_cache) > 100:
                oldest_key = min(self._self_rag_cache.items(), key=lambda x: x[1].get('_cached_at', 0))[0]
                del self._self_rag_cache[oldest_key]
            
            app_logger.info(f"[Self-RAG] 决策完成: '{query}' -> complexity={complexity}, need_retrieval={decision.get('need_retrieval')}")
            return decision
        except Exception as e:
            return {
                "need_retrieval": True,
                "confidence": 0.0,
                "reason": f"决策失败: {str(e)}",
                "complexity": complexity,
                "adaptive_top_k": default_top_k
            }

    def _vector_search(
        self,
        query: str,
        top_k: int,
        metadata_filter: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """执行向量检索"""
        # 确保集合存在
        try:
            self._ensure_collection_exists()
        except Exception as e:
            app_logger.error(f"确保集合存在失败: {e}")
            # 尝试恢复集合
            if not self.ensure_collection_ready():
                return []

        if not self._db_available or self.collection is None:
            return []

        try:
            self._refresh_vector_backend()
            if self.vector_backend is None:
                return []

            documents = []
            backend_results = self.vector_backend.search(
                VectorSearchRequest(
                    query=query,
                    top_k=top_k,
                    metadata_filter=metadata_filter,
                )
            )
            for result in backend_results:
                metadata = dict(result.metadata or {})

                # 检查源文件是否存在
                source_file = metadata.get('file_path', '')
                if source_file and not os.path.exists(source_file):
                    app_logger.warning(f"[向量检索] 跳过已删除的文件: {source_file}")
                    continue

                doc = Document(
                    page_content=result.content,
                    metadata={**metadata, "score": result.score}
                )
                documents.append(doc)

            return documents
        except Exception as e:
            app_logger.error(f"向量检索失败: {str(e)}")
            # 如果是集合不存在的错误，尝试恢复
            if "does not exist" in str(e):
                app_logger.warning(f"[向量检索] 集合不存在，尝试恢复...")
                if self.ensure_collection_ready():
                    # 重新尝试检索
                    try:
                        self._refresh_vector_backend()
                        if self.vector_backend is None:
                            return []
                        documents = []
                        backend_results = self.vector_backend.search(
                            VectorSearchRequest(
                                query=query,
                                top_k=top_k,
                                metadata_filter=metadata_filter,
                            )
                        )
                        for result in backend_results:
                            metadata = dict(result.metadata or {})
                            doc = Document(
                                page_content=result.content,
                                metadata={**metadata, "score": result.score}
                            )
                            documents.append(doc)
                        return documents
                    except Exception as retry_error:
                        app_logger.error(f"[向量检索] 重试失败: {retry_error}")
            return []

    def _keyword_search(self, query: str, top_k: int) -> List[Document]:
        """关键词匹配检索"""
        all_docs = self._vector_search("", 100)
        if not all_docs:
            return []

        query_lower = query.lower()
        scored_docs = []

        for doc in all_docs:
            content_lower = doc.page_content.lower()
            if query_lower in content_lower:
                score = content_lower.count(query_lower)
                scored_docs.append((score, doc))

        scored_docs.sort(key=lambda x: x[0], reverse=True)
        return [doc for _, doc in scored_docs[:top_k]]

    def _build_contextual_query(self, current_query: str, chat_history: List[Dict[str, str]], llm) -> str:
        """根据对话历史增强当前查询"""
        if not chat_history:
            return current_query
        
        app_logger.info(f"[查询增强] 开始处理，当前查询: '{current_query}', 历史轮数: {len(chat_history)}")

        # 构建对话历史文本（包含用户和助手）
        history_lines = []
        for msg in chat_history[-6:]:
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            if role == 'user':
                history_lines.append(f"用户: {content}")
            else:
                history_lines.append(f"助手: {content[:100]}...")  # 助手回复截断
        history_text = "\n".join(history_lines)
        
        app_logger.debug(f"[查询增强] 对话历史:\n{history_text}")

        contextual_prompt = f"""基于以下对话历史，将当前问题补充完整：

对话历史：
{history_text}

当前问题：{current_query}

返回JSON格式：
{{"completed_query": "补充完整的问题"}}
"""

        try:
            app_logger.info(f"[查询增强] 调用LLM进行查询补全...")
            response = llm.invoke(contextual_prompt)
            result = json.loads(response.content.strip())
            completed_query = result.get("completed_query", current_query)
            
            if completed_query != current_query:
                app_logger.info(f"[查询增强] 成功: '{current_query}' -> '{completed_query}'")
            else:
                app_logger.info(f"[查询增强] 无需增强，查询已完整")
            
            return completed_query
        except Exception as e:
            app_logger.warning(f"[查询增强] 失败: {e}，使用原查询")
            return current_query

    def retrieve_with_layers(
        self,
        query: str,
        top_k: int = 5,
        llm=None,
        use_hybrid: bool = True,
        use_rerank: bool = True,
        chat_history: List[Dict[str, str]] = None
    ) -> Dict[str, Any]:
        """
        三层架构检索 (Query Understanding -> Retrieval -> Generation Context)

        Args:
            query: 用户查询
            top_k: 基础 top_k
            llm: LLM 实例 (用于 Self-RAG 决策)
            use_hybrid: 是否使用混合检索
            use_rerank: 是否使用重排序
            chat_history: 对话历史

        Returns:
            包含 context (生成上下文), intent, documents 等
        """
        start_time = time.time()
        chat_history = chat_history or []

        # ========== Layer 1: Query Understanding ==========
        # 如果有 LLM，更新 query_layer
        if llm is not None:
            self.query_layer._llm = llm

        layer1_result = self.query_layer.process(query, chat_history)
        intent = layer1_result.intent
        entities = layer1_result.entities
        enhanced_query = layer1_result.enhanced_query
        complexity = layer1_result.complexity

        app_logger.info(f"[Layer 1] Query Understanding: intent={intent}, complexity={complexity}, entities={entities}")

        # ========== Layer 2: Retrieval ==========
        layer2_result = self.retrieval_layer.execute(
            query=enhanced_query,
            top_k=top_k,
            llm=llm,
            use_hybrid=use_hybrid,
            use_rerank=use_rerank,
            chat_history=chat_history,
            intent=intent,
            complexity=complexity
        )

        app_logger.info(f"[Layer 2] Retrieval: decision={layer2_result.decision}, "
                       f"docs={len(layer2_result.documents)}, "
                       f"hybrid={layer2_result.hybrid_applied}, "
                       f"rerank={layer2_result.rerank_applied}")

        # ========== Layer 3: Generation Context ==========
        layer3_result = self.generation_layer.build_context(
            documents=layer2_result.documents,
            intent=intent,
            query=query
        )

        app_logger.info(f"[Layer 3] Generation Context: context_length={len(layer3_result.context_text)}")

        return {
            "success": layer2_result.decision != "failed",
            "need_retrieval": layer2_result.decision == "retrieval",
            "context": layer3_result.context_text,
            "context_for_llm": self.generation_layer.format_for_llm(layer3_result),
            "intent": intent,
            "intent_confidence": layer1_result.intent_confidence,
            "entities": entities,
            "complexity": complexity,
            "documents": layer2_result.documents,
            "source_attributions": layer3_result.source_attributions,
            "decision_info": {
                "decision": layer2_result.decision,
                "complexity": layer2_result.complexity,
                "confidence": layer2_result.confidence,
                "reasoning": layer2_result.reasoning
            },
            "retrieval_time": time.time() - start_time,
            "adaptive_top_k": layer2_result.adaptive_top_k,
            "hybrid_applied": layer2_result.hybrid_applied,
            "rerank_applied": layer2_result.rerank_applied,
            "query": query,
            "enhanced_query": enhanced_query
        }

    def retrieve(self, query: str, top_k: int = 3, enable_self_rag: bool = True, llm=None,
                 use_cache: bool = True, use_hybrid: bool = True, use_rerank: bool = False,
                 chat_history: List[Dict[str, str]] = None,
                 retrieval_policy: Optional[Dict[str, Any]] = None,
                 trace_id: Optional[str] = None) -> Dict[str, Any]:
        """执行向量检索（带输入验证）"""
        # ========== 输入验证 ==========
        if not query or not isinstance(query, str):
            app_logger.warning("[retrieve] 无效输入: query为空或非字符串")
            return {
                "success": False,
                "error": "query不能为空",
                "documents": [],
                "decision_info": {}
            }

        query = query.strip()
        if len(query) > 4096:
            app_logger.warning(f"[retrieve] query过长: {len(query)}字符，截断至4096")
            query = query[:4096]

        if not isinstance(top_k, int) or top_k <= 0:
            app_logger.warning(f"[retrieve] 无效top_k: {top_k}，使用默认值3")
            top_k = 3
        elif top_k > 100:
            app_logger.warning(f"[retrieve] top_k过大: {top_k}，限制为100")
            top_k = 100

        if not isinstance(enable_self_rag, bool):
            app_logger.warning(f"[retrieve] enable_self_rag类型错误: {type(enable_self_rag)}，使用默认值True")
            enable_self_rag = True

        if chat_history is not None and not isinstance(chat_history, list):
            app_logger.warning(f"[retrieve] chat_history类型错误: {type(chat_history)}，设为空列表")
            chat_history = []
        # ========== 验证结束 ==========

        start_time = time.time()
        chat_history = chat_history or []
        effective_retrieval_policy = normalize_retrieval_policy(
            {
                **(retrieval_policy or {}),
                "enable_hybrid": use_hybrid,
                "enable_rerank": use_rerank,
            }
        )
        vector_metadata_filter = build_vector_metadata_filter(effective_retrieval_policy)
        
        # 调试日志：检查查询增强条件
        app_logger.info(f"[retrieve] Query: '{query}', chat_history: {len(chat_history)}, llm: {llm is not None}")
        if len(chat_history) > 0:
            app_logger.info(f"[retrieve] chat_history: {chat_history}")
        else:
            app_logger.warning(f"[retrieve] WARN: chat_history is empty, skip query enhancement")

        contextualized_query = query
        if len(chat_history) > 0 and llm is not None:
            app_logger.info(f"[retrieve] OK: Calling query enhancement...")
            contextualized_query = self._build_contextual_query(query, chat_history, llm)
            app_logger.info(f"[retrieve] Query enhancement done: '{query}' -> '{contextualized_query}'")
        else:
            if len(chat_history) == 0:
                app_logger.info(f"[retrieve] SKIP: chat_history is empty")
            if llm is None:
                app_logger.info(f"[retrieve] SKIP: llm is None")

        if use_cache:
            cached_result = self.cache.get(
                contextualized_query,
                top_k,
                enable_self_rag,
                retrieval_policy=effective_retrieval_policy,
            )
            if cached_result:
                self.metrics.record_cache_hit()
                cached_result = dict(cached_result)
                cached_result["cache_hit"] = True
                if trace_id:
                    cached_result["trace_id"] = trace_id
                return cached_result
            self.metrics.record_cache_miss()

        # 智能刷新：只在需要时刷新
        time_since_refresh = (datetime.now() - self.last_refresh_time).total_seconds()
        collection_count = self.collection.count() if self.collection else 0

        should_refresh = (
            self._need_initial_load or  # 首次加载
            (time_since_refresh > self._refresh_interval and collection_count == 0)  # 超时且集合为空
        )

        if should_refresh:
            if self._need_initial_load:
                app_logger.info(f"[refresh] 首次加载知识库...")
            elif collection_count == 0:
                app_logger.info(f"[refresh] 知识库为空，执行初始加载...")
            self.refresh_crag_knowledge_base()
            self._need_initial_load = False
        else:
            if collection_count > 0:
                app_logger.debug(f"[retrieve] 跳过刷新，集合已有 {collection_count} 个文档")

        actual_top_k = top_k
        decision_info = {}

        if enable_self_rag and llm:
            decision = self.self_rag_decision(contextualized_query, llm)
            decision_info = {
                "decision": decision,
                "complexity": decision.get("complexity", "unknown"),
                "confidence": decision.get("confidence", 0.0),
                "original_query": query,
                "contextualized_query": contextualized_query
            }

            if not decision["need_retrieval"]:
                result = {
                    "success": True,
                    "need_retrieval": False,
                    "reason": decision["reason"],
                    "documents": [],
                    "decision_info": decision_info,
                    "trace_id": trace_id,
                    "retrieval_policy": effective_retrieval_policy,
                }
                self.metrics.record_request(True, time.time() - start_time)
                self.metrics.record_self_rag_decision("no_retrieval")
                return result

            actual_top_k = decision.get("adaptive_top_k", top_k)
            self.metrics.record_self_rag_decision("retrieval")
        else:
            decision_info = {
                "original_query": query,
                "contextualized_query": contextualized_query
            }

        try:
            # 增加检索数量以给Rerank留出空间
            search_top_k = actual_top_k * 2 if use_rerank else actual_top_k
            
            # 执行向量检索
            docs = self._vector_search(
                contextualized_query,
                search_top_k,
                metadata_filter=vector_metadata_filter,
            )

            formatted_docs = []
            for doc in docs:
                formatted_docs.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "similarity_score": doc.metadata.get("score", 0.0),
                    "source_file": doc.metadata.get("file_path", "未知"),
                    "chunk_id": doc.metadata.get("chunk_id", "未知"),
                    "retrieval_method": "vector"  # 标记检索来源
                })
            
            # 混合检索：添加BM25结果
            hybrid_applied = False
            app_logger.info(f"[retrieve] use_hybrid={use_hybrid}, formatted_docs数量={len(formatted_docs)}")
            if use_hybrid:
                hybrid_start = time.time()

                # 获取所有文档用于BM25索引（只在文档数量变化时重建）
                if self.collection:
                    try:
                        current_count = self.collection.count()

                        # 只有在文档数量变化或首次时重建索引
                        if current_count != self._bm25_doc_count or not self._bm25_indexed:
                            all_docs = self.collection.get(include=["documents", "metadatas"])
                            bm25_docs = []
                            for i, doc_text in enumerate(all_docs.get("documents", [])):
                                metadata = all_docs.get("metadatas", [])[i] if i < len(all_docs.get("metadatas", [])) else {}
                                if not metadata_matches_filter(metadata, vector_metadata_filter):
                                    continue
                                bm25_docs.append({
                                    "content": doc_text,
                                    "metadata": metadata,
                                    "source_file": metadata.get("file_path", "未知"),
                                    "chunk_id": metadata.get("chunk_id", "未知")
                                })
                            app_logger.info(f"[BM25] 文档数量变化，重建索引，文档数: {len(bm25_docs)}")
                            self.bm25_retriever.index_documents(bm25_docs)
                            self._bm25_indexed = True
                            self._bm25_doc_count = current_count
                            # ProfessionalBM25 不提供 extracted_keywords 属性，使用文档数量作为统计
                            app_logger.info(f"[BM25] 索引已建立，文档数: {len(bm25_docs)}")
                        else:
                            app_logger.debug(f"[BM25] 使用缓存索引，文档数: {self._bm25_doc_count}")
                    except Exception as e:
                        app_logger.error(f"[BM25] 建立索引失败: {e}")
                        import traceback
                        traceback.print_exc()
                
                # 执行BM25检索
                bm25_results = self.bm25_retriever.search(contextualized_query, top_k=search_top_k)

                # 转换ProfessionalBM25结果为字典格式（含已删除文档过滤）
                bm25_dicts = []
                filtered_by_deletion = 0  # 统计被过滤的已删除文档数

                for r in bm25_results:
                    source_file = r.metadata.get("source_file", r.metadata.get("file_path", ""))
                    if not metadata_matches_filter(r.metadata, vector_metadata_filter):
                        continue

                    # 关键修复：检查源文件是否存在（与向量搜索保持一致，见1382-1386行）
                    if source_file and not os.path.exists(source_file):
                        app_logger.warning(f"[BM25检索] 跳过已删除的文件: {source_file}")
                        filtered_by_deletion += 1
                        continue

                    bm25_dicts.append({
                        "content": r.content,
                        "metadata": r.metadata,
                        "source_file": source_file or "未知",
                        "chunk_id": r.metadata.get("chunk_id", "未知"),
                        "bm25_score": r.score,
                        "doc_id": r.doc_id
                    })

                if filtered_by_deletion > 0:
                    app_logger.warning(f"[BM25检索] 已过滤 {filtered_by_deletion} 个来自已删除文件的文档块")

                if bm25_dicts:
                    if formatted_docs:
                        # 向量检索有结果，合并结果：RRF融合算法
                        k = 60  # RRF参数
                        SEMANTIC_WEIGHT = 1.5  # 语义搜索权重系数（1.5=推荐，语义搜索占比60%）
                        doc_scores = {}

                        app_logger.info(f"[混合检索] 语义搜索权重: {SEMANTIC_WEIGHT} (1.0=平衡, 1.5=语义优先, 2.0=强语义)")

                        # 向量检索结果打分（应用语义权重）
                        for rank, doc in enumerate(formatted_docs):
                            doc_key = (doc.get('source_file', ''), doc.get('chunk_id', ''))
                            doc_scores[doc_key] = doc_scores.get(doc_key, 0) + SEMANTIC_WEIGHT / (k + rank + 1)

                        # BM25结果打分（权重为1）
                        for rank, doc in enumerate(bm25_dicts):
                            doc_key = (doc.get('source_file', ''), doc.get('chunk_id', ''))
                            if doc_key in doc_scores:
                                doc_scores[doc_key] += 1 / (k + rank + 1)
                            else:
                                doc['retrieval_method'] = 'bm25'
                                doc['similarity_score'] = doc.get('bm25_score', 0)
                                formatted_docs.append(doc)
                                doc_scores[doc_key] = 1 / (k + rank + 1)

                        # 按融合分数排序
                        formatted_docs.sort(key=lambda x: doc_scores.get((x.get('source_file', ''), x.get('chunk_id', '')), 0), reverse=True)
                        formatted_docs = formatted_docs[:actual_top_k]
                        hybrid_applied = True

                        app_logger.info(f"[混合检索] BM25检索耗时: {time.time() - hybrid_start:.3f}s, 融合{len(bm25_dicts)}条结果")
                    else:
                        # 向量检索无结果，直接使用BM25结果
                        for doc in bm25_dicts:
                            doc['retrieval_method'] = 'bm25'
                            doc['similarity_score'] = doc.get('bm25_score', 0)
                        formatted_docs = bm25_dicts[:actual_top_k]
                        hybrid_applied = True

                        app_logger.info(f"[混合检索] 向量检索无结果，使用BM25结果: {len(formatted_docs)}条, 耗时: {time.time() - hybrid_start:.3f}s")

            # Rerank重排序
            rerank_applied = False
            if use_rerank and formatted_docs and self.reranker:
                rerank_start = time.time()
                formatted_docs = self.reranker.rerank(
                    contextualized_query,
                    formatted_docs,
                    top_k=actual_top_k
                )
                app_logger.info(f"Rerank耗时: {time.time() - rerank_start:.3f}s")
                rerank_applied = True

            retrieval_time = time.time() - start_time
            result = {
                "success": True,
                "need_retrieval": True,
                "documents": formatted_docs,
                "reason": "检索完成",
                "retrieval_time": retrieval_time,
                "actual_top_k": actual_top_k,
                "decision_info": decision_info,
                "query": query,
                "contextualized_query": contextualized_query,
                "retrieved_count": len(formatted_docs),
                "rerank_applied": rerank_applied,
                "hybrid_applied": hybrid_applied,  # 标记是否使用混合检索
                "trace_id": trace_id,
                "retrieval_policy": effective_retrieval_policy,
                "cache_hit": False,
            }

            self.metrics.record_request(True, retrieval_time)
            if use_cache:
                self.cache.set(
                    contextualized_query,
                    top_k,
                    enable_self_rag,
                    result,
                    retrieval_policy=effective_retrieval_policy,
                )
            app_logger.info(f"检索完成，返回{len(formatted_docs)}条文档，耗时{retrieval_time:.3f}s")
            
            # 记录检索到的具体文档信息
            if formatted_docs:
                app_logger.info(f"检索到的文档详情：")
                for i, doc in enumerate(formatted_docs, 1):
                    source_file = doc.get("source_file", "未知")
                    similarity = doc.get("similarity_score", 0.0)
                    content_preview = doc.get("content", "")[:100] + "..." if len(doc.get("content", "")) > 100 else doc.get("content", "")
                    app_logger.info(f"  文档{i}: 来源={source_file}, 相似度={similarity:.4f}, 内容预览={content_preview}")
            
            return result

        except Exception as e:
            app_logger.error(f"检索失败: {str(e)}")
            self.metrics.record_request(False, time.time() - start_time)
            return {
                "success": False,
                "error": str(e),
                "documents": [],
                "decision_info": decision_info,
                "trace_id": trace_id,
                "retrieval_policy": effective_retrieval_policy,
            }

    async def async_load_document(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """异步文档加载"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self.load_document, file_path, metadata)

    async def async_add_documents_to_vector_db(self, documents: List[Document]) -> List[str]:
        """异步文档入库"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self.add_documents_to_vector_db, documents)

    def run(self, query: str, top_k: int = 3, enable_self_rag: bool = True, llm=None,
            chat_history: List[Dict[str, str]] = None, use_hybrid: bool = True, use_rerank: bool = False) -> str:
        """RAG工具主入口"""
        retrieval_result = self.retrieve(
            query=query,
            top_k=top_k,
            enable_self_rag=enable_self_rag,
            llm=llm,
            use_cache=False,
            chat_history=chat_history,
            use_rerank=use_rerank
        )

        if not retrieval_result["success"]:
            return f"RAG检索失败: {retrieval_result.get('error', '未知错误')}"

        if not retrieval_result["need_retrieval"]:
            return f"无需检索: {retrieval_result['reason']}"

        if not retrieval_result["documents"]:
            return "未检索到相关文档"

        formatted_text = "[检索结果]\n"

        for i, doc in enumerate(retrieval_result["documents"], 1):
            formatted_text += f"\n{i}. 内容：{doc['content'][:200]}..."
            if doc['metadata']:
                formatted_text += f"\n   来源：{doc['metadata'].get('file_path', '未知')}"

        if "retrieval_time" in retrieval_result:
            formatted_text += f"\n\n检索耗时：{retrieval_result['retrieval_time']:.3f}s"

        return formatted_text

    def get_metrics(self) -> Dict[str, Any]:
        """获取系统指标"""
        return self.metrics.get_summary()

    def reset_metrics(self):
        """重置指标"""
        self.metrics.reset()

    def get_cache_stats(self) -> Dict[str, Any]:
        """获取缓存统计"""
        return self.cache.get_stats()

    def clear_cache(self):
        """手动清除缓存"""
        self.cache.clear()


rag_tool = RAGTool()
