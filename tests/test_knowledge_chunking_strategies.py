import importlib
import os
import sys
import tempfile
import types
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest import mock
from xml.sax.saxutils import escape

import docx

BACKEND_ROOT = Path(__file__).resolve().parents[1] / "backend"
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))


class _DummySentenceTransformer:
    def __init__(self, *args, **kwargs):
        self.args = args
        self.kwargs = kwargs

    def encode(self, texts, normalize_embeddings=False):
        return [[0.0] * 3 for _ in texts]


class _DummyCrossEncoder:
    def __init__(self, *args, **kwargs):
        self.args = args
        self.kwargs = kwargs

    def predict(self, pairs):
        return [0.0 for _ in pairs]


_sentence_transformers_stub = types.ModuleType("sentence_transformers")
_sentence_transformers_stub.SentenceTransformer = _DummySentenceTransformer
_sentence_transformers_stub.CrossEncoder = _DummyCrossEncoder

with mock.patch.dict(sys.modules, {"sentence_transformers": _sentence_transformers_stub}):
    rag_tool_module = importlib.import_module("tools.rag_tool")

Document = rag_tool_module.Document
RAGTool = rag_tool_module.RAGTool


def _xlsx_cell_reference(row_index, column_index):
    column_name = ""
    current = column_index
    while current:
        current, remainder = divmod(current - 1, 26)
        column_name = chr(65 + remainder) + column_name
    return f"{column_name}{row_index}"


def _xlsx_inline_cell_xml(row_index, column_index, value):
    cell_ref = _xlsx_cell_reference(row_index, column_index)
    if value in (None, ""):
        return f'<c r="{cell_ref}"/>'
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return f'<c r="{cell_ref}"><v>{value}</v></c>'
    text = escape(str(value))
    return (
        f'<c r="{cell_ref}" t="inlineStr"><is><t>{text}</t></is></c>'
    )


def _sheet_xml(rows):
    row_xml = []
    for row_index, row in enumerate(rows, start=1):
        cells = "".join(
            _xlsx_inline_cell_xml(row_index, column_index, value)
            for column_index, value in enumerate(row, start=1)
        )
        row_xml.append(f'<row r="{row_index}">{cells}</row>')
    sheet_data = "".join(row_xml)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f"<sheetData>{sheet_data}</sheetData>"
        "</worksheet>"
    )


def _write_xlsx(file_path, sheets):
    workbook_sheets = []
    workbook_relationships = []
    for index, (sheet_name, _) in enumerate(sheets, start=1):
        workbook_sheets.append(
            f'<sheet name="{escape(sheet_name)}" sheetId="{index}" '
            f'r:id="rId{index}"/>'
        )
        workbook_relationships.append(
            '<Relationship '
            f'Id="rId{index}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
            f'Target="worksheets/sheet{index}.xml"/>'
        )

    workbook_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook '
        'xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f"<sheets>{''.join(workbook_sheets)}</sheets>"
        "</workbook>"
    )
    workbook_rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        f"{''.join(workbook_relationships)}"
        "</Relationships>"
    )
    content_types_parts = [
        '<Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
    ]
    for index, _ in enumerate(sheets, start=1):
        content_types_parts.append(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        )
    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" '
        'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        f"{''.join(content_types_parts)}"
        "</Types>"
    )
    root_rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        "</Relationships>"
    )

    with zipfile.ZipFile(file_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", root_rels_xml)
        archive.writestr("xl/workbook.xml", workbook_xml)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels_xml)
        for index, (_, rows) in enumerate(sheets, start=1):
            archive.writestr(f"xl/worksheets/sheet{index}.xml", _sheet_xml(rows))


class KnowledgeChunkingStrategyTest(unittest.TestCase):
    def make_tool(self):
        tool = RAGTool.__new__(RAGTool)
        tool.text_splitter = SimpleNamespace(_chunk_size=400, _chunk_overlap=50)
        tool.metrics = SimpleNamespace(record_documents_processed=lambda count: None)
        return tool

    def test_detect_document_strategy_by_extension(self):
        tool = self.make_tool()

        self.assertEqual(
            tool._detect_document_strategy("sample.html"),
            "structured_document",
        )
        self.assertEqual(
            tool._detect_document_strategy("sample.docx"),
            "structured_document",
        )
        self.assertEqual(
            tool._detect_document_strategy("sample.xlsx"),
            "table_document",
        )
        self.assertEqual(
            tool._detect_document_strategy("sample.txt"),
            "recursive_text",
        )
        self.assertEqual(
            tool._detect_document_strategy("sample.xls"),
            "recursive_text",
        )

    def test_detect_document_strategy_for_pdf_with_outline(self):
        tool = self.make_tool()

        with mock.patch.object(tool, "_pdf_has_outline", return_value=True):
            self.assertEqual(
                tool._detect_document_strategy("sample.pdf"),
                "structured_document",
            )

    def test_load_document_sets_strategy_metadata_and_dispatches(self):
        tool = self.make_tool()

        temp_path = Path(__file__).resolve().parent / "_knowledge_chunking_test.pdf"
        temp_path.write_bytes(b"%PDF-1.4")

        try:
            def structured_loader(**kwargs):
                return [
                    Document(
                        page_content="stub",
                        metadata=dict(kwargs["metadata"]),
                    )
                ]

            with mock.patch.object(tool, "_pdf_has_outline", return_value=True):
                with mock.patch.object(
                    tool,
                    "_load_structured_document",
                    side_effect=structured_loader,
                ) as structured_mock:
                    documents = tool.load_document(
                        str(temp_path),
                        metadata={"source": "unit-test"},
                    )
        finally:
            temp_path.unlink(missing_ok=True)

        self.assertEqual(structured_mock.call_count, 1)
        self.assertEqual(len(documents), 1)
        self.assertEqual(
            documents[0].metadata["document_strategy"],
            "structured_document",
        )
        self.assertEqual(documents[0].metadata["source"], "unit-test")

    def test_recursive_splitter_preserves_punctuation_aware_separators(self):
        tool = self.make_tool()

        general_splitter = tool._build_recursive_splitter(
            chunk_size=400,
            chunk_overlap=50,
            content_type="general",
        )
        price_splitter = tool._build_recursive_splitter(
            chunk_size=400,
            chunk_overlap=50,
            content_type="price_list",
        )

        self.assertIn("\n\n", general_splitter._separators)
        self.assertIn("\n", general_splitter._separators)
        self.assertIn(". ", general_splitter._separators)
        self.assertIn("? ", general_splitter._separators)
        self.assertIn("! ", general_splitter._separators)
        self.assertIn(", ", general_splitter._separators)
        self.assertIn("; ", general_splitter._separators)
        self.assertIn(": ", general_splitter._separators)
        self.assertEqual(price_splitter._separators[0], "\n===== ")

    def test_html_structured_chunking_preserves_heading_path_and_strips_noise(self):
        tool = self.make_tool()

        tests_dir = Path(__file__).resolve().parent
        fd, temp_name = tempfile.mkstemp(
            suffix=".html",
            prefix="_knowledge_chunking_",
            dir=tests_dir,
        )
        os.close(fd)
        file_path = Path(temp_name)
        try:
            file_path.write_text(
                """
                <html>
                  <body>
                    <nav>Home Pricing Docs</nav>
                    <script>window.shouldNotAppear = true;</script>
                    <h1>Product Guide</h1>
                    <p>Overview for the supported workflow.</p>
                    <h2>Installation</h2>
                    <p>Install the package and configure the environment.</p>
                    <ul>
                      <li>Prepare tools</li>
                      <li>Run setup</li>
                    </ul>
                    <footer>Footer links should disappear.</footer>
                  </body>
                </html>
                """,
                encoding="utf-8",
            )

            documents = tool.load_document(
                str(file_path),
                chunk_size=500,
                chunk_overlap=0,
            )
        finally:
            file_path.unlink(missing_ok=True)

        joined_content = "\n\n".join(doc.page_content for doc in documents)

        self.assertIn("Product Guide", joined_content)
        self.assertIn("Product Guide > Installation", joined_content)
        self.assertIn("Overview for the supported workflow.", joined_content)
        self.assertIn("Prepare tools", joined_content)
        self.assertIn("Run setup", joined_content)
        self.assertNotIn("Home Pricing Docs", joined_content)
        self.assertNotIn("shouldNotAppear", joined_content)
        self.assertNotIn("Footer links should disappear.", joined_content)

    def test_docx_structured_chunking_preserves_heading_path(self):
        tool = self.make_tool()

        tests_dir = Path(__file__).resolve().parent
        fd, temp_name = tempfile.mkstemp(
            suffix=".docx",
            prefix="_knowledge_chunking_",
            dir=tests_dir,
        )
        os.close(fd)
        file_path = Path(temp_name)
        try:
            document = docx.Document()
            document.add_heading("Operations Manual", level=1)
            document.add_paragraph("General usage notes for operators.")
            document.add_heading("Safety", level=2)
            document.add_paragraph("Wear gloves during maintenance.")
            document.save(file_path)

            documents = tool.load_document(
                str(file_path),
                chunk_size=500,
                chunk_overlap=0,
            )
        finally:
            file_path.unlink(missing_ok=True)

        joined_content = "\n\n".join(doc.page_content for doc in documents)

        self.assertIn("Operations Manual", joined_content)
        self.assertIn("Operations Manual > Safety", joined_content)
        self.assertIn("General usage notes for operators.", joined_content)
        self.assertIn("Wear gloves during maintenance.", joined_content)

    def test_xlsx_chunking_keeps_sheet_headers_and_row_ranges(self):
        tool = self.make_tool()
        tests_dir = Path(__file__).resolve().parent
        fd, temp_name = tempfile.mkstemp(
            suffix=".xlsx",
            prefix="_knowledge_chunking_",
            dir=tests_dir,
        )
        os.close(fd)
        file_path = Path(temp_name)
        try:
            _write_xlsx(
                file_path,
                [
                    (
                        "Prices",
                        [
                            ["Product", "Price"],
                            ["Alpha", 99],
                            ["Beta", 199],
                            ["Gamma", 299],
                        ],
                    )
                ],
            )

            documents = tool.load_document(
                str(file_path),
                chunk_size=400,
                chunk_overlap=50,
            )
        finally:
            file_path.unlink(missing_ok=True)

        joined_content = "\n\n".join(doc.page_content for doc in documents)

        self.assertIn("Sheet: Prices", joined_content)
        self.assertIn("Headers: Product | Price", joined_content)
        self.assertIn("Rows 2-4", joined_content)
        self.assertTrue(all(doc.metadata["sheet_name"] == "Prices" for doc in documents))
        self.assertTrue(all(doc.metadata["row_range"] == "2-4" for doc in documents))

    def test_xlsx_chunking_falls_back_to_generated_headers(self):
        tool = self.make_tool()
        tests_dir = Path(__file__).resolve().parent
        fd, temp_name = tempfile.mkstemp(
            suffix=".xlsx",
            prefix="_knowledge_chunking_",
            dir=tests_dir,
        )
        os.close(fd)
        file_path = Path(temp_name)
        try:
            _write_xlsx(
                file_path,
                [
                    (
                        "Sheet1",
                        [
                            [None, ""],
                            ["Alpha", 99],
                        ],
                    )
                ],
            )

            documents = tool.load_document(
                str(file_path),
                chunk_size=400,
                chunk_overlap=50,
            )
        finally:
            file_path.unlink(missing_ok=True)

        joined_content = "\n\n".join(doc.page_content for doc in documents)

        self.assertIn("Headers: Column 1 | Column 2", joined_content)

    def test_read_document_content_supports_html_and_xlsx(self):
        from app.api.v1.knowledge_base import _read_document_content

        tests_dir = Path(__file__).resolve().parent
        html_fd, html_name = tempfile.mkstemp(
            suffix=".html",
            prefix="_knowledge_preview_",
            dir=tests_dir,
        )
        xlsx_fd, xlsx_name = tempfile.mkstemp(
            suffix=".xlsx",
            prefix="_knowledge_preview_",
            dir=tests_dir,
        )
        os.close(html_fd)
        os.close(xlsx_fd)
        html_path = Path(html_name)
        xlsx_path = Path(xlsx_name)
        try:
            html_path.write_text(
                "<html><body><h1>Guide</h1><p>Step one</p></body></html>",
                encoding="utf-8",
            )
            _write_xlsx(
                xlsx_path,
                [
                    (
                        "Prices",
                        [
                            ["Product", "Price"],
                            ["Alpha", 99],
                        ],
                    )
                ],
            )

            html_content = _read_document_content(html_path)
            xlsx_content = _read_document_content(xlsx_path)
        finally:
            html_path.unlink(missing_ok=True)
            xlsx_path.unlink(missing_ok=True)

        self.assertIn("Guide", html_content)
        self.assertIn("Step one", html_content)
        self.assertIn("[Prices]", xlsx_content)
        self.assertIn("Product | Price", xlsx_content)


if __name__ == "__main__":
    unittest.main()
